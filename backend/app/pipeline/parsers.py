"""Local document parsers: DOCX, PDF, XLSX."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass
class ParsedDoc:
    folder: str       # e.g. "APM", "AWP", "Samples"
    filename: str     # original basename, e.g. "Lumina Grand_2. APM (8 Mar) (V3).docx"
    text: str         # extracted text, lightly Markdown-ish


def parse_docx(path: Path) -> str:
    """Extract paragraphs and tables from a .docx as plain text / light Markdown."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        parts.append("")  # blank line before table
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n".join(parts).strip()


def parse_pdf(path: Path) -> str:
    """Extract text from a .pdf using pdfplumber, one block per page."""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Page {i} ---\n{text.strip()}")
    return "\n\n".join(pages).strip()


def parse_xlsx(path: Path) -> str:
    """Extract each sheet of a .xlsx as a Markdown-ish table."""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        rows_iter = sheet.iter_rows(values_only=True)
        header = next(rows_iter, None)
        if header is None:
            continue
        header_cells = [str(c) if c is not None else "" for c in header]
        parts.append("| " + " | ".join(header_cells) + " |")
        parts.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for row in rows_iter:
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                parts.append("| " + " | ".join(cells) + " |")
        parts.append("")
    return "\n".join(parts).strip()


PARSERS_BY_EXT = {
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".xlsx": parse_xlsx,
}

ARTEFACT_FOLDERS = [
    "APM", "AWP", "Guidelines",
    "Process SOP", "Process Understanding", "Samples",
]


def parse_folder(project_root: Path, skip_filenames: set[str] | None = None) -> list[ParsedDoc]:
    """Parse every known-extension file in each artefact folder.

    Files whose basename appears in `skip_filenames` are still discovered
    but not parsed or returned (they are held out from downstream stages).
    Unknown extensions are silently skipped. Parse failures are re-raised
    with the offending path for the caller to log.
    """
    skip = skip_filenames or set()
    out: list[ParsedDoc] = []
    for folder in ARTEFACT_FOLDERS:
        fdir = project_root / folder
        if not fdir.is_dir():
            continue
        for p in sorted(fdir.iterdir()):
            if not p.is_file():
                continue
            if p.name.startswith("~$"):  # Word/Excel lock artefacts
                continue
            if p.name in skip:
                continue
            parser = PARSERS_BY_EXT.get(p.suffix.lower())
            if parser is None:
                continue
            try:
                text = parser(p)
            except Exception as e:
                raise RuntimeError(f"failed to parse {p}: {e}") from e
            out.append(ParsedDoc(folder=folder, filename=p.name, text=text))
    return out


def persist_parsed(parsed: list[ParsedDoc], out_root: Path) -> None:
    """Write each ParsedDoc to `<out_root>/parsed/<folder>/<basename>.md`."""
    for d in parsed:
        target_dir = out_root / "parsed" / d.folder
        target_dir.mkdir(parents=True, exist_ok=True)
        stem = Path(d.filename).stem
        (target_dir / f"{stem}.md").write_text(d.text, encoding="utf-8")
