"""DOCX renderer -- emits Sample-style 5-column issue table + exception tables.

Applies a style_spec (from app.pipeline.prompts.styling) to match template appearance
(landscape A4, palette, fonts, paragraph spacing).

Layout constants (column widths, header/footer content structure, cell margins,
header/footer distance) were extracted once from the approved past-audit issue
log and frozen here -- the renderer does not read that file at runtime.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Mm, Pt

from .prompts.styling import DEFAULT_STYLE_SPEC, normalise_style_spec

ISSUE_HEADERS = ["S/N", "Findings", "Possible Impact", "Recommendations", "Comments"]
_RISK_BANNER = {
    "High": "HIGH RISK",
    "Medium": "MEDIUM RISK",
    "Low": "LOW RISK",
}

# A4 in millimetres (portrait). Landscape swaps width/height.
_A4_SHORT_MM = 210
_A4_LONG_MM = 297

# --- Layout constants frozen from the approved past-audit issue log. ---
# Column widths in twips (1/20 pt). Must sum to the printable page width.
SUMMARY_COL_TWIPS = [565, 9595, 1080, 1080, 1091, 909]   # S/N | Title | H | M | L | Page
ISSUE_COL_TWIPS = [552, 5120, 2410, 2838, 3039]          # S/N | Findings | Impact | Recs | Comments
# Cell margins: 108 twips (~0.075") on L/R matches Word default, which issue/
# exception tables use implicitly. Samples uses similar per-cell margins (55/52)
# -- anything non-zero keeps text from touching the border.
SUMMARY_CELL_MAR_TWIPS = {"left": 108, "right": 108}
# Header/footer distance is now read from style_spec.page (Guideline §18/§19).


def _set_cell_shading(cell, hex_color: str | None) -> None:
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.replace("#", "").upper())


def _apply_paragraph_spacing(cell, spacing: dict[str, Any]) -> None:
    before = Pt(spacing["before_pt"])
    after = Pt(spacing["after_pt"])
    line_min = Pt(spacing["line_at_least_pt"])
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = before
        pf.space_after = after
        pf.line_spacing = line_min


def _set_cell_vmerge(cell, val: str | None) -> None:
    """val='restart' starts a vertical merge; val=None (or no value) continues one.

    Pass 'restart' on the top cell and None on cells below that should merge up.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vMerge"))
    if existing is not None:
        tcPr.remove(existing)
    vm = OxmlElement("w:vMerge")
    if val:
        vm.set(qn("w:val"), val)
    tcPr.append(vm)


def _set_row_height(row, twips: int, rule: str = "atLeast") -> None:
    """Set explicit row height in twips. rule='atLeast' allows growth;
    'exact' pins height."""
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn("w:trHeight"))
    if existing is not None:
        trPr.remove(existing)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(twips)))
    if rule:
        h.set(qn("w:hRule"), rule)
    trPr.append(h)


def _set_cell_valign(cell, val: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcPr.append(v)


def _set_paragraph_align(para, align: str) -> None:
    """align: 'left' | 'center' | 'right'."""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = mapping[align]


def _set_cell_paragraphs_align(cell, align: str) -> None:
    for para in cell.paragraphs:
        _set_paragraph_align(para, align)


def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    font_family: str = "Arial",
    size_pt: int = 10,
    spacing: dict[str, Any] | None = None,
    align: str | None = None,
    valign: str | None = None,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    if text:
        for i, line in enumerate(text.split("\n")):
            run = para.add_run(line) if i == 0 else cell.add_paragraph().add_run(line)
            run.font.name = font_family
            run.font.size = Pt(size_pt)
            run.font.bold = bold
    if spacing:
        _apply_paragraph_spacing(cell, spacing)
    if align:
        _set_cell_paragraphs_align(cell, align)
    if valign:
        _set_cell_valign(cell, valign)


def _merge_row(row) -> None:
    first = row.cells[0]
    for c in row.cells[1:]:
        first.merge(c)


def _category_code(category: str) -> tuple[str, str]:
    """Split 'A. PERSONAL DATA PROTECTION...' -> ('A.', 'PERSONAL DATA PROTECTION...')."""
    category = (category or "").strip()
    if not category:
        return ("", "")
    parts = category.split(None, 1)
    if len(parts) == 2 and parts[0].rstrip(".").isalpha() and len(parts[0].rstrip(".")) <= 2:
        return (parts[0], parts[1])
    return ("", category)


def _set_table_column_widths(table, twips_per_col: list[int]) -> None:
    """Set tblGrid + per-cell width to lock column widths."""
    tbl = table._element
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in twips_per_col:
        g = OxmlElement("w:gridCol")
        g.set(qn("w:w"), str(int(w)))
        grid.append(g)
    # Insert tblGrid right after tblPr.
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(grid)
    else:
        tbl.insert(0, grid)
    # Force table layout = fixed so cell widths are honored.
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")
    # Set width on every cell in every row.
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(twips_per_col):
                cell.width = Emu(int(twips_per_col[ci]) * 635)


def _set_table_cell_margins(table, mar_twips: dict[str, int]) -> None:
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in mar_twips.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


def _apply_page_setup(doc: Document, page: dict[str, Any]) -> None:
    section = doc.sections[0]
    # Force landscape A4 per user mandate; page spec overrides only margins.
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Mm(_A4_LONG_MM)
    section.page_height = Mm(_A4_SHORT_MM)
    margins = page.get("margin_inches", {})
    if "left" in margins:
        section.left_margin = Inches(margins["left"])
        section.right_margin = Inches(margins["right"])
        section.top_margin = Inches(margins["top"])
        section.bottom_margin = Inches(margins["bottom"])
    # Guideline §18/§19: header and footer positioned 0.49" from page edge.
    section.header_distance = Inches(page.get("header_distance_inches", 0.49))
    section.footer_distance = Inches(page.get("footer_distance_inches", 0.49))


def _clear_paragraph(para) -> None:
    for run in list(para.runs):
        run._element.getparent().remove(run._element)


def _add_header(section, entity_legal: str, project_name: str, fiscal_year: str) -> None:
    """Three-line centered header per Samples convention. All caps (Guideline §18)."""
    header = section.header
    # Clear existing paragraphs.
    for p in header.paragraphs:
        _clear_paragraph(p)
    # Ensure 3 paragraphs for 3 lines (first para already exists).
    paragraphs = list(header.paragraphs)
    while len(paragraphs) < 3:
        paragraphs.append(header.add_paragraph())

    lines = [
        ("CITY DEVELOPMENTS LIMITED", True),
        (f"{fiscal_year} INTERNAL AUDIT REPORT" if fiscal_year else "INTERNAL AUDIT REPORT", False),
        (
            f"{entity_legal.upper()} (\u201c{project_name.upper()}\u201d)" if project_name
            else entity_legal.upper(),
            False,
        ),
    ]
    for para, (text, bold) in zip(paragraphs, lines):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Match Samples: before=after=60 twips (3pt) on each content line,
        # single line spacing. Gives tight inter-line spacing and a small
        # breath below the entity name before the body starts.
        pf = para.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = bold

    # Trailing empty paragraph below the entity name (Samples convention)
    # so the header has a small gap before the body content begins.
    trailer = header.add_paragraph()
    tf = trailer.paragraph_format
    tf.space_before = Pt(0)
    tf.space_after = Pt(0)
    tf.line_spacing = 1.0


def _add_footer(section, page_width_emu: int, left_margin_emu: int, right_margin_emu: int) -> None:
    """Footer layout per Samples:
        [PAGE] | AUDIT FINDINGS  <right-tab>  CONFIDENTIAL
    PAGE + " | AUDIT FINDINGS" bold Arial 10, left-aligned.
    CONFIDENTIAL bold Arial 10, grey 7F7F7F, right-aligned via tab stop.
    Paragraph has a top border (single, 4pt, D9D9D9).
    """
    footer = section.footer
    for p in footer.paragraphs:
        _clear_paragraph(p)
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Remove inherited pStyle ("Footer") so its default tab stops don't interfere
    # with our custom right-aligned tab.
    printable_emu = page_width_emu - left_margin_emu - right_margin_emu
    pPr = para._element.get_or_add_pPr()
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pPr.remove(pStyle)
    old_tabs = pPr.find(qn("w:tabs"))
    if old_tabs is not None:
        pPr.remove(old_tabs)
    tabs = OxmlElement("w:tabs")
    # Clear any inherited Footer-style tabs (defensive; common positions).
    for clear_pos in ("4680", "9360"):
        clr = OxmlElement("w:tab")
        clr.set(qn("w:val"), "clear")
        clr.set(qn("w:pos"), clear_pos)
        tabs.append(clr)
    right_twip = int(printable_emu / 635)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(right_twip))
    tabs.append(tab)
    pPr.append(tabs)

    # Top border.
    old_bdr = pPr.find(qn("w:pBdr"))
    if old_bdr is not None:
        pPr.remove(old_bdr)
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "D9D9D9")
    pBdr.append(top)
    pPr.append(pBdr)

    def _styled_run(text: str, *, color_hex: str | None = None, bold: bool = True) -> None:
        r = para.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = bold
        if color_hex:
            # Set run color via OOXML (python-docx RGBColor avoided to keep string).
            rPr = r._element.get_or_add_rPr()
            old_col = rPr.find(qn("w:color"))
            if old_col is not None:
                rPr.remove(old_col)
            col = OxmlElement("w:color")
            col.set(qn("w:val"), color_hex)
            rPr.append(col)

    # PAGE field as bold Arial 10.
    def _wrap_field_char(ftype: str) -> None:
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), "Arial")
        rf.set(qn("w:hAnsi"), "Arial")
        rpr.append(rf)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")
        rpr.append(sz)
        b = OxmlElement("w:b")
        rpr.append(b)
        r.append(rpr)
        if ftype == "instr":
            instr = OxmlElement("w:instrText")
            instr.text = " PAGE   \\* MERGEFORMAT "
            r.append(instr)
        else:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), ftype)
            r.append(fc)
        para._element.append(r)

    _wrap_field_char("begin")
    _wrap_field_char("instr")
    _wrap_field_char("separate")
    _wrap_field_char("end")

    _styled_run(" | AUDIT FINDINGS")
    _styled_run("\t")
    _styled_run("CONFIDENTIAL", color_hex="7F7F7F")


def _add_summary_index(doc: Document, draft: list[dict[str, Any]], style: dict[str, Any]) -> None:
    if not draft:
        return

    body_font = style["fonts"]["body"]
    body_spacing = style["paragraph_spacing"]["body"]
    colors = style["colors"]
    summary_bg = colors["summary_header_bg"]
    category_bg = colors["category_banner_bg"]
    marker = colors["risk_level_marker"]

    tbl = doc.add_table(rows=0, cols=6)
    tbl.style = "Table Grid"
    _set_table_cell_margins(tbl, SUMMARY_CELL_MAR_TWIPS)

    fam = body_font["family"]
    sz = body_font["size_pt"]

    # Header row 1: Issues (2x2 merge) | Audit Risk Level (1x3 merge) | Page (2x1 merge)
    # Merge BEFORE writing to avoid python-docx preserving empty paragraphs.
    h1 = tbl.add_row()
    h1.cells[0].merge(h1.cells[1])
    h1.cells[2].merge(h1.cells[3])
    h1.cells[2].merge(h1.cells[4])
    _write_cell(h1.cells[0], "Issues", bold=True, font_family=fam, size_pt=sz,
                spacing=body_spacing, align="center", valign="center")
    _write_cell(h1.cells[2], "Audit Risk Level", bold=True, font_family=fam, size_pt=sz,
                spacing=body_spacing, align="center", valign="center")
    _write_cell(h1.cells[5], "Page", bold=True, font_family=fam, size_pt=sz,
                spacing=body_spacing, align="center", valign="center")
    for c in h1.cells:
        _set_cell_shading(c, summary_bg)
    # Start vertical merge on the Issues block (C0+C1) and Page (C5).
    _set_cell_vmerge(h1.cells[0], "restart")
    _set_cell_vmerge(h1.cells[5], "restart")

    # Header row 2: Issues (vMerge continue) | High | Medium | Low | Page (vMerge continue)
    h2 = tbl.add_row()
    h2.cells[0].merge(h2.cells[1])
    # Leave the merged Issues cell empty — it will visually merge with R0's "Issues".
    for i, label in [(2, "High"), (3, "Medium"), (4, "Low")]:
        _write_cell(h2.cells[i], label, bold=True, font_family=fam, size_pt=sz,
                    spacing=body_spacing, align="center", valign="center")
    for c in h2.cells:
        _set_cell_shading(c, summary_bg)
    # Continue vertical merge on Issues (C0+C1) and Page (C5).
    _set_cell_vmerge(h2.cells[0], None)
    _set_cell_vmerge(h2.cells[5], None)

    # Group by category preserving first-seen order.
    seen_categories: list[str] = []
    grouped: dict[str, list[dict[str, Any]]] = {}
    for issue in draft:
        key = issue.get("category", "") or ""
        if key not in grouped:
            grouped[key] = []
            seen_categories.append(key)
        grouped[key].append(issue)

    for cat in seen_categories:
        code, name = _category_code(cat)
        cat_row = tbl.add_row()
        # Merge C1..C5 first, then write.
        cat_row.cells[1].merge(cat_row.cells[2])
        cat_row.cells[1].merge(cat_row.cells[3])
        cat_row.cells[1].merge(cat_row.cells[4])
        cat_row.cells[1].merge(cat_row.cells[5])
        _write_cell(cat_row.cells[0], code, bold=True, font_family=fam, size_pt=sz,
                    spacing=body_spacing, align="center", valign="center")
        _write_cell(cat_row.cells[1], name, bold=True, font_family=fam, size_pt=sz,
                    spacing=body_spacing, align="left", valign="center")
        for cell in (cat_row.cells[0], cat_row.cells[1]):
            _set_cell_shading(cell, category_bg)

        for issue in grouped[cat]:
            iid_short = issue.get("id", "").replace("I-", "A")
            title = issue.get("title", "")
            risk = issue.get("risk_level", "") or ""
            row = tbl.add_row()
            _write_cell(row.cells[0], iid_short, bold=True, font_family=fam, size_pt=sz,
                        spacing=body_spacing, align="center", valign="center")
            _write_cell(row.cells[1], title, font_family=fam, size_pt=sz,
                        spacing=body_spacing, align="left", valign="center")
            # Risk columns: colored cell with no text (visual marker only).
            for ci in (2, 3, 4):
                _write_cell(row.cells[ci], "", spacing=body_spacing, valign="center")
            if risk == "High":
                _set_cell_shading(row.cells[2], marker.get("High"))
            elif risk == "Medium":
                _set_cell_shading(row.cells[3], marker.get("Medium"))
            elif risk == "Low":
                _set_cell_shading(row.cells[4], marker.get("Low"))
            _write_cell(row.cells[5], "", font_family=fam, size_pt=sz,
                        spacing=body_spacing, align="center", valign="center")

    _set_table_column_widths(tbl, SUMMARY_COL_TWIPS)


def _add_issue_table(doc: Document, issue: dict[str, Any], style: dict[str, Any]) -> None:
    iid_short = issue.get("id", "I-?").replace("I-", "A")
    title = issue.get("title", "")
    risk_level = issue.get("risk_level", "") or ""
    category = issue.get("category", "") or ""

    body_font = style["fonts"]["body"]
    body_spacing = style["paragraph_spacing"]["body"]
    colors = style["colors"]
    fam = body_font["family"]
    sz = body_font["size_pt"]

    table = doc.add_table(rows=0, cols=5)
    table.style = "Table Grid"

    # Header row (S/N | Findings | Possible Impact | Recommendations | Comments)
    hdr = table.add_row()
    for i, h in enumerate(ISSUE_HEADERS):
        _write_cell(hdr.cells[i], h, bold=True,
                    font_family=fam, size_pt=sz, spacing=body_spacing,
                    align="center", valign="center")
        _set_cell_shading(hdr.cells[i], colors["issue_table_header_bg"])

    # Category banner row (merged across all columns) — merge first, then write.
    if category:
        cat_row = table.add_row()
        _merge_row(cat_row)
        _write_cell(cat_row.cells[0], category, bold=True,
                    font_family=fam, size_pt=sz, spacing=body_spacing,
                    align="left", valign="center")
        _set_cell_shading(cat_row.cells[0], colors["category_banner_bg"])

    # Risk banner row (merged across all columns) — merge first, then write.
    banner = _RISK_BANNER.get(risk_level, "")
    if banner:
        risk_row = table.add_row()
        _merge_row(risk_row)
        _write_cell(risk_row.cells[0], banner, bold=True,
                    font_family=fam, size_pt=sz, spacing=body_spacing,
                    align="left", valign="center")
        _set_cell_shading(risk_row.cells[0], colors["risk_banner_bg"])

    # Title row: S/N | title spanning 4 cols — merge first, then write.
    title_row = table.add_row()
    title_row.cells[1].merge(title_row.cells[2])
    title_row.cells[1].merge(title_row.cells[3])
    title_row.cells[1].merge(title_row.cells[4])
    _write_cell(title_row.cells[0], iid_short, bold=True,
                font_family=fam, size_pt=sz, spacing=body_spacing,
                align="center", valign="center")
    _write_cell(title_row.cells[1], title, bold=True,
                font_family=fam, size_pt=sz, spacing=body_spacing,
                align="left", valign="center")

    # Body row
    body_row = table.add_row()
    _write_cell(body_row.cells[0], "", spacing=body_spacing,
                align="center", valign="center")

    finding_parts = [issue.get("finding", "").strip()]
    rc = issue.get("root_cause", "").strip()
    if rc:
        finding_parts.append(f"Root Cause: {rc}")
    th = issue.get("theme", "").strip()
    if th:
        finding_parts.append(f"Theme: {th}")
    _write_cell(body_row.cells[1], "\n\n".join(p for p in finding_parts if p),
                font_family=fam, size_pt=sz, spacing=body_spacing, align="left")

    impact_parts = []
    ri = issue.get("risk_impact", "").strip()
    if ri:
        impact_parts.append(f"Risk Impact:\n{ri}")
    fi = issue.get("financial_impact", "").strip()
    if fi:
        impact_parts.append(f"Financial Impact:\n{fi}")
    _write_cell(body_row.cells[2], "\n\n".join(impact_parts),
                font_family=fam, size_pt=sz, spacing=body_spacing, align="left")

    _write_cell(body_row.cells[3], issue.get("recommendation", "").strip(),
                font_family=fam, size_pt=sz, spacing=body_spacing, align="left")

    comment_parts = []
    mc = issue.get("management_comments", "").strip()
    if mc:
        comment_parts.append(f"Management's Comments: {mc}")
    ap = issue.get("action_plan", "").strip()
    if ap:
        comment_parts.append(f"Action Plan:\n{ap}")
    resp = issue.get("responsibility", "").strip()
    if resp:
        comment_parts.append(f"Responsibility:\n{resp}")
    td = issue.get("target_date", "").strip()
    if td:
        comment_parts.append(f"Target Date:\n{td}")
    _write_cell(body_row.cells[4], "\n\n".join(comment_parts),
                font_family=fam, size_pt=sz, spacing=body_spacing, align="left")

    _set_table_column_widths(table, ISSUE_COL_TWIPS)


def _exception_col_widths(n_cols: int, printable_twips: int) -> list[int]:
    """Frozen from Samples: first col (S/N) ~550 twips; remaining share printable."""
    if n_cols <= 0:
        return []
    sn = 550 if n_cols > 1 else printable_twips
    remaining = printable_twips - sn
    if n_cols == 1:
        return [printable_twips]
    per = remaining // (n_cols - 1)
    widths = [sn] + [per] * (n_cols - 1)
    # Absorb rounding error into the last column so totals equal printable_twips.
    widths[-1] += printable_twips - sum(widths)
    return widths


def _add_page_break_paragraph(doc: Document) -> None:
    """Insert a standalone paragraph whose only run is <w:br w:type='page'/>.

    Kept for callers that explicitly need a plain page break; the pipeline now
    uses `_add_break(doc, style)` which honours style_spec.page_break_type.
    """
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_section_break_next_page(doc: Document) -> None:
    """Insert a 'Section Break (Next Page)' per Guideline §9.

    python-docx's Document.add_section(WD_SECTION.NEW_PAGE) appends a fresh
    section with a nextPage start-type. Headers/footers inherit from the
    previous section by default, so our header/footer only need to be authored
    once (on the first section).
    """
    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_break(doc: Document, style: dict[str, Any]) -> None:
    """Insert the break type requested by style_spec.page.page_break_type."""
    kind = style.get("page", {}).get("page_break_type", "section_next_page")
    if kind == "section_next_page":
        _add_section_break_next_page(doc)
    else:
        _add_page_break_paragraph(doc)


def _add_exception_table(
    doc: Document, exc: dict[str, Any], style: dict[str, Any],
    *, printable_twips: int,
) -> None:
    title = exc.get("title", "").strip()
    headers = exc.get("headers") or []
    rows = exc.get("rows") or []
    if not headers or not rows:
        return

    exc_font = style["fonts"]["exception"]
    exc_spacing = style["paragraph_spacing"]["exception"]
    fam = exc_font["family"]
    sz = exc_font["size_pt"]

    # Page-break so the description table starts on its own page (Samples convention).
    _add_break(doc, style)

    if title:
        p = doc.add_paragraph()
        # Guideline §11: Before 6pt, After 6pt, At least 13pt between title
        # and the table. Body paragraph_spacing carries those values.
        body_spacing = style["paragraph_spacing"]["body"]
        pf = p.paragraph_format
        pf.space_before = Pt(body_spacing["before_pt"])
        pf.space_after = Pt(body_spacing["after_pt"])
        pf.line_spacing = Pt(body_spacing["line_at_least_pt"])
        run = p.add_run(title)
        run.font.name = fam
        run.font.size = Pt(sz)
        run.font.bold = True
        run.font.underline = True

    tbl = doc.add_table(rows=0, cols=len(headers))
    tbl.style = "Table Grid"

    # Samples uses grey D9D9D9 for description-table header (not the issue-table yellow).
    header_fill = style["colors"].get("category_banner_bg") or "D9D9D9"
    hdr = tbl.add_row()
    for i, h in enumerate(headers):
        _write_cell(hdr.cells[i], str(h), bold=True,
                    font_family=fam, size_pt=sz, spacing=exc_spacing,
                    align="center", valign="center")
        _set_cell_shading(hdr.cells[i], header_fill)

    # Guideline §11 alignment: "left" (Centre Left) for text, "center" for
    # dates / non-monetary numbers / S/N, "right" (Centre Right) for monetary
    # amounts. The drafting LLM emits column_aligns; fall back to a heuristic
    # (S/N centered, everything else left) when absent or malformed.
    first_header = (str(headers[0]).strip().lower() if headers else "")
    sn_col_is_first = first_header in {"s/n", "sn", "no.", "no", "#"}
    col_aligns_raw = exc.get("column_aligns") or []
    valid = {"left", "center", "right"}
    if (isinstance(col_aligns_raw, list)
            and len(col_aligns_raw) == len(headers)
            and all(a in valid for a in col_aligns_raw)):
        col_aligns = list(col_aligns_raw)
    else:
        col_aligns = [
            "center" if (i == 0 and sn_col_is_first) else "left"
            for i in range(len(headers))
        ]

    for r in rows:
        row = tbl.add_row()
        cells = list(r) + [""] * (len(headers) - len(r))
        for i, val in enumerate(cells[: len(headers)]):
            _write_cell(row.cells[i], str(val),
                        font_family=fam, size_pt=sz, spacing=exc_spacing,
                        align=col_aligns[i], valign="center")

    _set_table_column_widths(tbl, _exception_col_widths(len(headers), printable_twips))


def _add_evidence_list(doc: Document, refs: list[str], style: dict[str, Any]) -> None:
    if not refs:
        return
    exc_font = style["fonts"]["exception"]
    # Spacer paragraph so "Evidence references:" is not flush against the table above.
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Evidence references:")
    run.font.name = exc_font["family"]
    run.font.size = Pt(exc_font["size_pt"])
    run.font.bold = True
    run.font.italic = True
    for ref in refs:
        line = doc.add_paragraph(style="List Bullet")
        run = line.runs[0] if line.runs else line.add_run()
        run.font.name = exc_font["family"]
        run.font.size = Pt(exc_font["size_pt"])
        run.text = str(ref)


def _add_review_procedures_table(
    doc: Document,
    procedures: list[dict[str, Any]],
    style: dict[str, Any],
    *,
    printable_twips: int,
) -> None:
    """Appendix bảng 'Internal Audit Review Procedures' — 3 columns (Scope |
    Key Processes | Summarised Work Program). Mirrors Samples T7:
      R0: merged title row (summary_header_bg, centered, bold)
      R1: merged empty spacer row
      R2: column headers (category_banner_bg, centered, bold)
      R3+: data rows; Scope column vMerged across consecutive rows with same value.
    """
    if not procedures:
        return

    body_font = style["fonts"]["body"]
    body_spacing = style["paragraph_spacing"]["body"]
    colors = style["colors"]
    summary_bg = colors["summary_header_bg"]
    category_bg = colors["category_banner_bg"]
    fam = body_font["family"]
    sz = body_font["size_pt"]

    _add_break(doc, style)

    tbl = doc.add_table(rows=0, cols=3)
    tbl.style = "Table Grid"

    # R0: title row — content row, Guideline §2 body spacing (6/6/13) applies.
    # trHeight uses 'atLeast' so row can grow if font metrics push higher.
    r0 = tbl.add_row()
    r0.cells[0].merge(r0.cells[1])
    r0.cells[0].merge(r0.cells[2])
    _write_cell(r0.cells[0], "Internal Audit Review Procedures", bold=True,
                font_family=fam, size_pt=sz, spacing=body_spacing,
                align="center", valign="center")
    _set_cell_shading(r0.cells[0], summary_bg)
    _set_row_height(r0, 680, rule="atLeast")

    # R1: thin empty spacer row (Samples trHeight=43 twips, ~2pt). Being a
    # visual spacer with no content, §2 body spacing is intentionally skipped
    # so it stays as a hairline between title and header rows.
    r1 = tbl.add_row()
    r1.cells[0].merge(r1.cells[1])
    r1.cells[0].merge(r1.cells[2])
    _write_cell(r1.cells[0], "", valign="center")
    _set_row_height(r1, 43, rule="exact")

    # R2: column headers — content row, body spacing applies.
    r2 = tbl.add_row()
    for i, label in enumerate(("Scope", "Key Processes", "Summarised Work Program")):
        _write_cell(r2.cells[i], label, bold=True,
                    font_family=fam, size_pt=sz, spacing=body_spacing,
                    align="center", valign="center")
        _set_cell_shading(r2.cells[i], category_bg)
    _set_row_height(r2, 567, rule="atLeast")

    # Data rows. Apply vMerge on Scope column for consecutive duplicates.
    prev_scope: str | None = None
    for proc in procedures:
        scope = (proc.get("scope") or "").strip()
        kp = (proc.get("key_process") or "").strip()
        wp = (proc.get("work_program") or "").strip()
        row = tbl.add_row()
        # Write scope text only on the first row of a new group; continue merge otherwise.
        if scope and scope == prev_scope:
            _write_cell(row.cells[0], "", font_family=fam, size_pt=sz,
                        spacing=body_spacing, align="center", valign="center")
            _set_cell_vmerge(row.cells[0], None)
        else:
            _write_cell(row.cells[0], scope, bold=True, font_family=fam, size_pt=sz,
                        spacing=body_spacing, align="center", valign="center")
            _set_cell_vmerge(row.cells[0], "restart")
        _write_cell(row.cells[1], kp, font_family=fam, size_pt=sz,
                    spacing=body_spacing, align="left", valign="center")
        _write_cell(row.cells[2], wp, font_family=fam, size_pt=sz,
                    spacing=body_spacing, align="left", valign="center")
        prev_scope = scope

    # Column widths frozen from Samples T7 (2684 / 3260 / 7826), rescaled to
    # the current printable width so totals match the page.
    ref = [2684, 3260, 7826]
    total = sum(ref)
    widths = [int(w * printable_twips / total) for w in ref]
    widths[-1] += printable_twips - sum(widths)
    _set_table_column_widths(tbl, widths)


def render(
    draft: list[dict[str, Any]],
    project_name: str,
    version: str,
    out_path: Path,
    *,
    style_spec: dict[str, Any] | None = None,
    entity_legal: str = "",
    fiscal_year: str = "",
    review_procedures: list[dict[str, Any]] | None = None,
) -> None:
    style = normalise_style_spec(style_spec) if style_spec is not None else DEFAULT_STYLE_SPEC

    doc = Document()
    _apply_page_setup(doc, style["page"])

    # Header + footer per Samples convention (values extracted as constants above).
    section = doc.sections[0]
    _add_header(section, entity_legal=entity_legal or project_name,
                project_name=project_name, fiscal_year=fiscal_year)
    _add_footer(
        section,
        page_width_emu=int(section.page_width),
        left_margin_emu=int(section.left_margin),
        right_margin_emu=int(section.right_margin),
    )

    body_font = style["fonts"]["body"]
    normal = doc.styles["Normal"]
    normal.font.name = body_font["family"]
    normal.font.size = Pt(body_font["size_pt"])

    _add_summary_index(doc, draft, style)

    # Printable width in twips, for proportional description-table column widths.
    printable_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    printable_twips = int(printable_emu / 635)

    for issue in draft:
        # Each issue table starts on its own page (Samples convention).
        _add_break(doc, style)
        _add_issue_table(doc, issue, style)
        exc = issue.get("exceptions")
        if isinstance(exc, dict):
            _add_exception_table(doc, exc, style, printable_twips=printable_twips)
        _add_evidence_list(doc, issue.get("evidence_refs") or [], style)

    if review_procedures:
        _add_review_procedures_table(
            doc, review_procedures, style, printable_twips=printable_twips,
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
