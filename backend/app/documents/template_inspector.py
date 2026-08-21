"""Extract styling signals (palette, fonts, page setup) from an existing DOCX.

Used to feed real template values into the LLM styling step and to provide
deterministic fallbacks for the renderer.
"""
from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn

_CATEGORY_CODE_RE = re.compile(r"^[A-Z]\.$")
_RISK_BANNER_RE = re.compile(r"^(HIGH|MEDIUM|LOW)\s+RISK$", re.I)
_ISSUE_HEADER_TOKENS = {"s/n", "findings", "possible impact", "recommendations", "comments"}
_SUMMARY_HEADER_TOKENS = {"issues", "audit risk level", "page", "high", "medium", "low"}


def _cell_fill(cell) -> str | None:
    shd = cell._tc.find(f".//{qn('w:shd')}")
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill == "auto":
        return None
    return fill.upper()


def _dominant_font(doc: Document) -> tuple[str, int]:
    families: Counter[str] = Counter()
    sizes: Counter[int] = Counter()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                families[run.font.name] += 1
            if run.font.size:
                sizes[run.font.size.pt] += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            families[run.font.name] += 1
                        if run.font.size:
                            sizes[int(run.font.size.pt)] += 1
    fam = families.most_common(1)[0][0] if families else "Arial"
    size = sizes.most_common(1)[0][0] if sizes else 10
    return fam, size


def inspect_templates(paths: list[Path]) -> dict[str, Any]:
    """Analyse multiple templates and merge role votes across them.

    The first path's page setup and dominant font win; `roled_colors` votes
    are combined across all paths so a file with more risk-level examples
    can enrich the palette.
    """
    analyses = [inspect_template(p) for p in paths if p.is_file()]
    if not analyses:
        return {}
    base = analyses[0]
    merged: dict[str, Counter[str]] = {}
    for a in analyses:
        for role, color in (a.get("roled_colors") or {}).items():
            merged.setdefault(role, Counter())[color] += 1
    base["roled_colors"] = {r: ctr.most_common(1)[0][0] for r, ctr in merged.items()}
    base["sources"] = [a.get("source") for a in analyses]
    return base


def inspect_template(path: Path) -> dict[str, Any]:
    """Extract a compact styling profile from a DOCX template."""
    doc = Document(str(path))
    section = doc.sections[0]
    orientation = "landscape" if section.orientation == WD_ORIENTATION.LANDSCAPE else "portrait"

    family, size = _dominant_font(doc)

    # Collect unique fill colors grouped by rough role based on text content.
    fills: dict[str, list[str]] = {}
    role_votes: dict[str, Counter[str]] = {}

    def _vote(role: str, color: str) -> None:
        role_votes.setdefault(role, Counter())[color] += 1

    for ti, table in enumerate(doc.tables):
        rows = list(table.rows)
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                fill = _cell_fill(cell)
                if not fill:
                    continue
                txt = cell.text.strip()
                txt_lc = txt.lower()
                fills.setdefault(fill, []).append(f"T{ti}:{txt[:40]!r}")

                if txt_lc in _SUMMARY_HEADER_TOKENS:
                    _vote("summary_header_bg", fill)
                if txt_lc in _ISSUE_HEADER_TOKENS:
                    _vote("issue_table_header_bg", fill)
                if _CATEGORY_CODE_RE.match(txt):
                    _vote("category_banner_bg", fill)
                if _RISK_BANNER_RE.match(txt):
                    _vote("risk_banner_bg", fill)
                # Risk markers: empty cell under H/M/L header column in summary table.
                # Identify by checking the header row's column text.
                if not txt and ri >= 2 and ci in (2, 3, 4):
                    # Find the header row (first row with "High"/"Medium"/"Low")
                    header_map: dict[int, str] = {}
                    for hrow in rows[:3]:
                        for hi, hcell in enumerate(hrow.cells):
                            ht = hcell.text.strip().lower()
                            if ht in ("high", "medium", "low"):
                                header_map[hi] = ht.capitalize()
                    if ci in header_map:
                        _vote(f"risk_level_marker.{header_map[ci]}", fill)

    roled_colors = {role: ctr.most_common(1)[0][0] for role, ctr in role_votes.items()}

    return {
        "source": path.name,
        "page": {
            "orientation": orientation,
            "width_emu": section.page_width,
            "height_emu": section.page_height,
            "margin_emu": {
                "left": section.left_margin,
                "right": section.right_margin,
                "top": section.top_margin,
                "bottom": section.bottom_margin,
            },
            "margin_inches": round(section.left_margin / 914400, 2) if section.left_margin else None,
        },
        "dominant_font": {"family": family, "size_pt": size},
        "fill_colors": fills,
        "roled_colors": roled_colors,
    }
