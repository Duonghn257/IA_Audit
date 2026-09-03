from __future__ import annotations

from io import BytesIO
from pathlib import Path

from app.application.audit_pipeline import PipelineRequest, PipelineResult
from app.bootstrap.api import create_app
from docx import Document
from fastapi.testclient import TestClient

from tests.integration.test_discovery_api import _create_project, _settings


def _central_docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _upload_central_knowledge(client: TestClient) -> None:
    guideline = client.post(
        "/api/v1/central-knowledge/guidelines",
        files={
            "file": (
                "writing-guideline.docx",
                _central_docx_bytes("Central guideline original"),
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        },
    )
    assert guideline.status_code == 200, guideline.text
    template = client.put(
        "/api/v1/central-knowledge/template",
        files={
            "file": (
                "template.docx",
                _central_docx_bytes("Central template"),
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        },
    )
    assert template.status_code == 200, template.text


class SuccessfulAuditPipeline:
    def __init__(self) -> None:
        self.requests: list[PipelineRequest] = []

    def run(self, request: PipelineRequest, *, reporter=None) -> PipelineResult:
        self.requests.append(request)
        assert request.auditor_input is not None
        assert request.issues_path is None
        assert request.version == "v0.1"
        assert (request.project_path / "AWP" / "scope.docx").is_file()
        assert (request.project_path / "APM" / "risk.docx").is_file()
        assert (
            request.project_path
            / "Process Understanding"
            / "evidence.docx"
        ).is_file()
        assert (
            request.project_path / "Process SOP" / "criteria.docx"
        ).is_file()
        assert (
            request.project_path / "Samples" / "approved-report.docx"
        ).is_file()
        guideline_path = (
            request.project_path / "Guidelines" / "writing-guideline.docx"
        )
        assert guideline_path.is_file()
        assert Document(guideline_path).paragraphs[0].text == (
            "Central guideline original"
        )
        assert (request.project_path / "Output" / "template.docx").is_file()
        if reporter is not None:
            from app.application.audit_pipeline import PipelineProgress

            reporter(PipelineProgress("RENDERING", "DOCX generated", 8))
        assert request.run_directory is not None
        request.run_directory.mkdir(parents=True)
        output = request.run_directory / "Audit Test_Issue Log v0.1.docx"
        document = Document()
        document.add_heading("Generated audit issue log")
        document.add_paragraph(request.auditor_input[0]["observed_gap"])
        document.save(output)
        return PipelineResult(
            version="v0.1",
            run_directory=request.run_directory,
            output_path=output,
            issue_count=len(request.auditor_input),
        )


class FlakyAuditPipeline(SuccessfulAuditPipeline):
    def __init__(self) -> None:
        super().__init__()
        self.attempts = 0

    def run(self, request: PipelineRequest, *, reporter=None) -> PipelineResult:
        self.attempts += 1
        if self.attempts == 1:
            raise RuntimeError("Temporary audit pipeline failure")
        return super().run(request, reporter=reporter)


def _create_manual_candidate(
    client: TestClient, project_id: str, version_id: str
) -> dict[str, object]:
    response = client.post(
        f"/api/v1/projects/{project_id}/versions/{version_id}/issues",
        json={
            "title_hint": "Access review evidence should be retained",
            "observed_gap": "Quarterly access review evidence was incomplete.",
            "risk_category": "Operational",
            "status": "READY_FOR_REVIEW",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()


def test_audit_requires_candidates_in_selected_version(tmp_path: Path) -> None:
    app = create_app(settings=_settings(tmp_path))
    with TestClient(app) as client:
        project_id, version_id = _create_project(client)
        response = client.post(
            f"/api/v1/projects/{project_id}/versions/{version_id}/audit-jobs",
            json={"issue_revision": 0},
        )
        assert response.status_code == 422, response.text
        assert response.json()["error"]["code"] == "AUDIT_PREFLIGHT_FAILED"
        assert "candidate issues" in response.json()["error"]["message"]
    app.state.database.dispose()


def test_audit_requires_central_knowledge(tmp_path: Path) -> None:
    app = create_app(settings=_settings(tmp_path))
    with TestClient(app) as client:
        project_id, version_id = _create_project(client)
        _create_manual_candidate(client, project_id, version_id)
        version = client.get(
            f"/api/v1/projects/{project_id}/versions/{version_id}"
        ).json()
        response = client.post(
            f"/api/v1/projects/{project_id}/versions/{version_id}/audit-jobs",
            json={"issue_revision": version["issue_revision"]},
        )
        assert response.status_code == 422
        assert response.json()["error"]["code"] == "AUDIT_PREFLIGHT_FAILED"
        assert "Central knowledge is not ready" in (
            response.json()["error"]["message"]
        )
    app.state.database.dispose()


def test_audit_freezes_database_candidates_and_publishes_docx(
    tmp_path: Path,
) -> None:
    pipeline = SuccessfulAuditPipeline()
    app = create_app(settings=_settings(tmp_path), audit_pipeline=pipeline)
    with TestClient(app) as client:
        project_id, version_id = _create_project(client)
        _upload_central_knowledge(client)
        candidate = _create_manual_candidate(client, project_id, version_id)
        version = client.get(
            f"/api/v1/projects/{project_id}/versions/{version_id}"
        ).json()
        started = client.post(
            f"/api/v1/projects/{project_id}/versions/{version_id}/audit-jobs",
            json={"issue_revision": version["issue_revision"]},
        )
        assert started.status_code == 202, started.text
        job_id = started.json()["job_id"]
        job = client.get(f"/api/v1/jobs/{job_id}").json()
        assert job["state"] == "SUCCEEDED"
        assert job["attempt_count"] == 1
        assert pipeline.requests[0].auditor_input == [
            {
                "issue_id": candidate["issue_id"],
                "origin": "MANUAL",
                "status": "READY_FOR_REVIEW",
                "title_hint": candidate["title_hint"],
                "observed_gap": candidate["observed_gap"],
                "evidence_summary": None,
                "evidence_refs": [],
                "sop_refs": [],
                "risk_category": "Operational",
                "confidence": None,
                "validation_flags": [],
                "source_refs": [],
            }
        ]
        outputs = client.get(
            f"/api/v1/projects/{project_id}/versions/{version_id}/outputs"
        )
        assert outputs.status_code == 200, outputs.text
        assert len(outputs.json()) == 1
        output = outputs.json()[0]
        assert output["status"] == "CURRENT"
        downloaded = client.get(output["download_url"])
        assert downloaded.status_code == 200, downloaded.text
        assert downloaded.content[:2] == b"PK"
        events = client.get(f"/api/v1/jobs/{job_id}/events").json()
        assert [event["stage"] for event in events] == [
            "PREPARING_SOURCE",
            "RENDERING",
            "PUBLISHING",
        ]
        stream = client.get(f"/api/v1/jobs/{job_id}/events/stream")
        assert stream.status_code == 200
        assert "event: progress" in stream.text
        assert "event: end" in stream.text
    app.state.database.dispose()


def test_audit_retry_reuses_frozen_input_snapshot(tmp_path: Path) -> None:
    pipeline = FlakyAuditPipeline()
    app = create_app(settings=_settings(tmp_path), audit_pipeline=pipeline)
    with TestClient(app) as client:
        project_id, version_id = _create_project(client)
        _upload_central_knowledge(client)
        candidate = _create_manual_candidate(client, project_id, version_id)
        version = client.get(
            f"/api/v1/projects/{project_id}/versions/{version_id}"
        ).json()
        started = client.post(
            f"/api/v1/projects/{project_id}/versions/{version_id}/audit-jobs",
            json={"issue_revision": version["issue_revision"]},
        )
        job_id = started.json()["job_id"]
        assert client.get(f"/api/v1/jobs/{job_id}").json()["state"] == "FAILED"
        overwritten = client.post(
            "/api/v1/central-knowledge/guidelines",
            files={
                "file": (
                    "writing-guideline.docx",
                    _central_docx_bytes("Central guideline overwritten"),
                    (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                )
            },
        )
        assert overwritten.status_code == 200, overwritten.text
        updated = client.put(
            f"/api/v1/projects/{project_id}/versions/{version_id}/issues/"
            f"{candidate['issue_id']}",
            json={
                "row_version": candidate["row_version"],
                "observed_gap": "This edit happened after the job was frozen.",
                "risk_category": "Operational",
                "status": "READY_FOR_REVIEW",
            },
        )
        assert updated.status_code == 200, updated.text
        retried = client.post(
            f"/api/v1/jobs/{job_id}/retry",
            json={"reason": "Pipeline dependency recovered"},
        )
        assert retried.status_code == 202, retried.text
        succeeded = client.get(f"/api/v1/jobs/{job_id}").json()
        assert succeeded["state"] == "SUCCEEDED"
        assert succeeded["attempt_count"] == 2
        assert pipeline.requests[0].auditor_input[0]["observed_gap"] == (
            "Quarterly access review evidence was incomplete."
        )
    app.state.database.dispose()
