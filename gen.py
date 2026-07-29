from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

output_path = Path("./data/NovaBot_Phase2_Software_Requirements_Specification_v0.9.docx")

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

# -----------------------
# Styles
# -----------------------
styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(4)
styles["Normal"].paragraph_format.line_spacing = 1.05

style_specs = [
    ("Title", "Aptos Display", 25, "0F172A", True),
    ("Subtitle", "Aptos", 13, "0F766E", False),
    ("Heading 1", "Aptos Display", 17, "0F4E4A", True),
    ("Heading 2", "Aptos Display", 13.5, "0F172A", True),
    ("Heading 3", "Aptos", 11.2, "0F766E", True),
]
for name, font_name, size, color, bold in style_specs:
    st = styles[name]
    st.font.name = font_name
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = bold

styles["Heading 1"].paragraph_format.space_before = Pt(12)
styles["Heading 1"].paragraph_format.space_after = Pt(6)
styles["Heading 2"].paragraph_format.space_before = Pt(9)
styles["Heading 2"].paragraph_format.space_after = Pt(4)
styles["Heading 3"].paragraph_format.space_before = Pt(6)
styles["Heading 3"].paragraph_format.space_after = Pt(3)

if "Requirement" not in styles:
    st = styles.add_style("Requirement", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Aptos"
    st.font.size = Pt(9.7)
    st.paragraph_format.left_indent = Inches(0.15)
    st.paragraph_format.right_indent = Inches(0.08)
    st.paragraph_format.space_before = Pt(2)
    st.paragraph_format.space_after = Pt(4)

if "Callout" not in styles:
    st = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Aptos"
    st.font.size = Pt(10)
    st.font.color.rgb = RGBColor.from_string("0F4E4A")
    st.paragraph_format.left_indent = Inches(0.16)
    st.paragraph_format.right_indent = Inches(0.16)
    st.paragraph_format.space_before = Pt(4)
    st.paragraph_format.space_after = Pt(7)

# -----------------------
# Helpers
# -----------------------
def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def style_table(table, header_fill="0F766E", header_color="FFFFFF", first_col_bold=False, font_size=8.5):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for run in p.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(font_size)
            if i == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor.from_string(header_color)
                        run.font.bold = True
            elif i % 2 == 0:
                shade_cell(cell, "F8FAFC")
            if first_col_bold and i > 0 and j == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string("0F4E4A")

def add_bullets(items, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)

def add_callout(text, fill="CCFBF1", border="0F766E"):
    p = doc.add_paragraph(style="Callout")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), border)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.bold = True
    return p

def add_requirement(req_id, title, statement, priority="Must", status="Draft", note=None):
    p = doc.add_paragraph(style="Requirement")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F8FAFC")
    pPr.append(shd)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "CBD5E1")
    border.append(bottom)
    pPr.append(border)
    r = p.add_run(f"{req_id} — {title}\n")
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0F4E4A")
    p.add_run(statement)
    r2 = p.add_run(f"\nPriority: {priority} | Status: {status}")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string("64748B")
    if note:
        r3 = p.add_run(f" | Note: {note}")
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = RGBColor.from_string("B45309")

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_section_header(title, text=None):
    doc.add_heading(title, level=1)
    if text:
        doc.add_paragraph(text)

# Header/footer
header = section.header
hp = header.paragraphs[0]
hp.text = "NOVABOT PHASE 2 | SOFTWARE REQUIREMENTS SPECIFICATION | DRAFT v0.9"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.name = "Aptos"
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string("64748B")

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = fp.add_run("NovaBot Phase 2 SRS • Confidential Draft • Trang ")
rr.font.name = "Aptos"
rr.font.size = Pt(7.5)
rr.font.color.rgb = RGBColor.from_string("64748B")
add_page_number(fp)

# -----------------------
# Cover
# -----------------------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("NOVABOT PHASE 2")
r.font.name = "Aptos Display"
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string("0F766E")

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Software Requirements Specification")

p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("System-level requirements for Business → Store → Fanpage architecture")

doc.add_paragraph("")
meta = doc.add_table(rows=6, cols=2)
meta_data = [
    ("Document ID", "NVB-SRS-P2"),
    ("Version", "0.9 Draft"),
    ("Status", "For Review / Confirmation"),
    ("Product Owner", "Nguyễn Hoàng Dương"),
    ("Project", "NovaBot Phase 2"),
    ("Baseline", "NovaBot Phase 2 — Thiết kế cập nhật v2.0"),
]
for i, (k, v) in enumerate(meta_data):
    meta.cell(i, 0).text = k
    meta.cell(i, 1).text = v
style_table(meta, header_fill="FFFFFF", header_color="0F172A", first_col_bold=True, font_size=10)
for row in meta.rows:
    shade_cell(row.cells[0], "F1F5F9")

doc.add_paragraph("")
add_callout(
    "Đây là SRS cấp hệ thống. Tài liệu xác định phạm vi, kiến trúc, yêu cầu chức năng dùng chung, yêu cầu phi chức năng và tích hợp. "
    "Chi tiết feature và Acceptance Criteria của từng màn sẽ nằm trong Screen Specification/Jira Story."
)
doc.add_page_break()

# -----------------------
# Document Control
# -----------------------
doc.add_heading("Document Control", level=1)

doc.add_heading("Revision History", level=2)
t = doc.add_table(rows=1, cols=5)
for i, h in enumerate(["Version", "Date", "Author", "Change Summary", "Status"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("0.9", "15/07/2026", "Nguyễn Hoàng Dương / OpenAI-assisted draft", "Initial system-level SRS based on confirmed Phase 2 design", "For Review"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

doc.add_heading("Approval Matrix", level=2)
t = doc.add_table(rows=1, cols=4)
for i, h in enumerate(["Role", "Reviewer", "Responsibility", "Approval"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("Product Owner", "Nguyễn Hoàng Dương", "Scope, business rules, priorities", "Pending"),
    ("Design", "TBD", "Information architecture and UX feasibility", "Pending"),
    ("Frontend Lead", "TBD", "Navigation, state and frontend feasibility", "Pending"),
    ("Backend/AI Lead", "TBD", "Data, API, security and AI feasibility", "Pending"),
    ("QA", "TBD", "Testability and release criteria", "Pending"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

doc.add_heading("Requirement Status Legend", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Status", "Meaning", "Action"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("Confirmed", "Đã được chốt trong tài liệu thiết kế hoặc bởi Product Owner", "Có thể dùng làm baseline"),
    ("Draft", "Được đề xuất dựa trên thiết kế hiện tại", "Cần team review"),
    ("TBD", "Thiếu thông tin hoặc cần quyết định", "Product Owner phải xác nhận"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

doc.add_page_break()

# -----------------------
# Static TOC
# -----------------------
doc.add_heading("Mục lục", level=1)
toc = [
    "1. Introduction",
    "2. Product Overview",
    "3. User Classes, Roles and Permission Model",
    "4. System Context and Information Architecture",
    "5. Functional Requirements",
    "6. Data and Business Rules",
    "7. External Interface and Integration Requirements",
    "8. Non-functional Requirements",
    "9. Error Handling and Common UI States",
    "10. Security, Privacy and Authorization",
    "11. Logging, Monitoring and Auditability",
    "12. Deployment, Migration and Backward Compatibility",
    "13. Requirement Traceability and Change Control",
    "14. Assumptions and Constraints",
    "15. Open Questions / Confirmations Required",
    "Appendix A. Route Map",
    "Appendix B. Requirement ID Convention",
    "Appendix C. System-level Release Acceptance",
]
for item in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item)
doc.add_page_break()

# -----------------------
# 1. Introduction
# -----------------------
add_section_header("1. Introduction")

doc.add_heading("1.1 Purpose", level=2)
doc.add_paragraph(
    "Tài liệu này mô tả yêu cầu phần mềm cấp hệ thống cho NovaBot Phase 2. "
    "Mục tiêu là tạo baseline chung cho Product, Design, Frontend, Backend/AI, QA và DevOps trước khi đi vào Screen Specification, Jira Story và implementation."
)

doc.add_heading("1.2 Product Scope", level=2)
add_bullets([
    "Quản trị một Business/account có một hoặc nhiều Store.",
    "Mỗi Store quản lý dữ liệu bán hàng, chính sách, cấu hình AI, fanpage, hội thoại và deal riêng.",
    "Một Store có thể liên kết nhiều Fanpage.",
    "Fanpage là filter/breakdown trong Store, không phải cấp navigation thứ ba.",
    "Hỗ trợ vận hành AI bán hàng và chăm sóc khách hàng qua hội thoại.",
    "Cung cấp Dashboard, AI Quality, Usage & Billing, thành viên và phân quyền.",
])

doc.add_heading("1.3 Out of Scope for This SRS", level=2)
add_bullets([
    "Chi tiết pixel, màu sắc, spacing và component cụ thể của từng màn.",
    "Acceptance Criteria chi tiết cho từng feature.",
    "Thiết kế database vật lý và implementation detail.",
    "Kế hoạch thương mại, giá bán cụ thể và hợp đồng khách hàng.",
    "Tích hợp kênh ngoài phạm vi được xác nhận cho Phase 2.",
])

doc.add_heading("1.4 Referenced Documents", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Document", "Purpose", "Status"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("NovaBot Phase 2 — Thiết kế cập nhật v2.0", "Product architecture, dual sidebar, route map and module scope", "Confirmed baseline"),
    ("NovaBot Phase 1 Audit & Migration Plan", "Screen inventory and migration matrix", "Working document"),
    ("Screen Specifications", "Per-screen features, states and acceptance criteria", "To be created incrementally"),
    ("Jira Backlog", "Execution, estimate and sprint status", "To be created from approved specifications"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

doc.add_heading("1.5 Definitions and Abbreviations", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Term", "Definition", "Status"]):
    t.rows[0].cells[i].text = h
terms = [
    ("Business", "Phạm vi quản trị cao nhất của một doanh nghiệp/account, chứa Store, thành viên, quota và billing.", "Confirmed"),
    ("Store", "Đơn vị vận hành bán hàng, chứa sản phẩm, chính sách, AI config, fanpage, messages và deals.", "Confirmed"),
    ("Fanpage", "Kênh bán hàng thuộc Store; là nguồn hội thoại/deal và filter phân tích.", "Confirmed"),
    ("Pending human", "Hội thoại đang cần người thật tiếp quản nhưng chưa hoàn thành xử lý.", "Draft — cần rule chi tiết"),
    ("Human takeover", "Trạng thái hoặc hành động người thật tiếp quản hội thoại từ AI.", "Draft — cần rule chi tiết"),
    ("Deal detected", "Cơ hội mua hàng do AI hoặc hệ thống nhận diện từ hội thoại.", "Draft — cần tiêu chí"),
    ("Fallback", "AI không đủ thông tin hoặc không đủ tự tin để trả lời theo logic được cấu hình.", "Draft — cần ngưỡng/quy tắc"),
    ("Usage", "Tài nguyên tính theo gói như AI replies, messages, connected pages hoặc members.", "Draft — cần pricing model"),
]
for rd in terms:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.2)

# -----------------------
# 2. Product Overview
# -----------------------
add_section_header("2. Product Overview")

doc.add_heading("2.1 Current Baseline", level=2)
doc.add_paragraph(
    "Phase 1 đã có các luồng cơ bản: đăng nhập/đăng ký, Dashboard cơ bản, Store, Fanpage, Sản phẩm, Chính sách, Tin nhắn, Trợ lý AI và bot AI xử lý hội thoại. "
    "Một số KPI Dashboard vẫn là mock data và navigation hiện tại chưa phân tách rõ Business/Store context."
)

doc.add_heading("2.2 Phase 2 Product Direction", level=2)
add_callout("Business level → Store level → Fanpage filter / breakdown")
add_bullets([
    "Business quản trị tập trung: Dashboard tổng, Members & Permissions, Usage & Billing, Business Settings.",
    "Store vận hành: Dashboard, Messages, Deals, Products, Policies, Promotions, Pages, AI Assistant, AI Quality, Store Settings.",
    "Fanpage không có sidebar riêng; dùng để lọc hoặc drill-down dữ liệu.",
    "Dashboard chuyển từ ecommerce dashboard sang AI Sales Operation Dashboard.",
    "Light theme và semantic theme token là baseline.",
])

doc.add_heading("2.3 Primary Product Objectives", level=2)
add_bullets([
    "Giúp Owner biết Store nào đang cần xử lý.",
    "Giúp nhân viên nhanh chóng nhìn thấy hội thoại Pending human và Deal cần follow-up.",
    "Giúp Store cải thiện chất lượng AI từ hội thoại thật.",
    "Đảm bảo dữ liệu, quyền và usage được tách đúng Business/Store/Fanpage.",
    "Giảm phụ thuộc vào người vận hành sản phẩm bằng tài liệu và requirement có thể truy vết.",
])

doc.add_heading("2.4 High-level Module Map", level=2)
t = doc.add_table(rows=1, cols=4)
for i, h in enumerate(["Scope", "Module", "Purpose", "Phase 2"]):
    t.rows[0].cells[i].text = h
module_rows = [
    ("Auth", "Login / Register / Recovery", "Xác thực và khôi phục tài khoản", "In scope"),
    ("Onboarding", "Create Store → Connect Pages → Import Products → Policies → Assistant → Test Bot", "Thiết lập Store đầu tiên", "In scope, detail TBD"),
    ("Business", "Business Dashboard", "Tổng quan nhiều Store", "In scope"),
    ("Business", "Members & Permissions", "Quản trị tập trung thành viên và quyền theo Store", "In scope"),
    ("Business", "Usage & Billing", "Gói, quota, hóa đơn và breakdown theo Store", "In scope"),
    ("Store", "Store Dashboard", "Tổng quan vận hành AI với 4 tab", "In scope"),
    ("Store", "Messages", "Inbox đa fanpage trong Store", "In scope"),
    ("Store", "Deals", "Workspace xử lý cơ hội bán hàng", "In scope"),
    ("Store", "Products / Policies / Promotions", "Dữ liệu bán hàng và knowledge", "In scope"),
    ("Store", "Pages", "Quản lý fanpage liên kết", "In scope"),
    ("Store", "AI Assistant / AI Quality", "Cấu hình, test và cải thiện AI", "In scope"),
    ("Store", "Store Settings", "Cấu hình Store", "In scope"),
]
for rd in module_rows:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.0)

# -----------------------
# 3. User Roles
# -----------------------
add_section_header("3. User Classes, Roles and Permission Model")

doc.add_heading("3.1 Proposed Role Model", level=2)
add_callout(
    "TBD quan trọng: tài liệu trước đây có lúc sử dụng 3 role Owner/Admin/Staff; thiết kế Phase 2 có thêm Business Admin, Store Manager, Sales và Support. "
    "Cần Product Owner chốt role model chính thức trước khi triển khai permission."
, fill="FEF3C7", border="D97706")

t = doc.add_table(rows=1, cols=4)
for i, h in enumerate(["Role", "Scope", "Typical Capability", "Status"]):
    t.rows[0].cells[i].text = h
role_rows = [
    ("Business Owner", "Business", "Toàn quyền quản trị, Store, members, billing và settings", "Draft"),
    ("Business Admin", "Business / assigned Stores", "Quản trị theo permission được cấp", "Draft"),
    ("Store Manager", "Assigned Store(s)", "Quản lý vận hành, products, policies, pages, messages", "Draft"),
    ("Sales", "Assigned Store(s)", "Messages, deals, follow-up và customer actions", "Draft"),
    ("Support", "Assigned Store(s)", "Messages, support cases, có thể hạn chế deal/settings", "Draft"),
]
for rd in role_rows:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.3)

doc.add_heading("3.2 Permission Principles", level=2)
add_requirement("FR-PERM-001", "Backend authorization", "Mọi API nhạy cảm phải kiểm tra authorization ở Backend; không được chỉ dựa vào việc ẩn nút trên Frontend.", "Must", "Draft")
add_requirement("FR-PERM-002", "Business scope", "Quyền Business-level phải áp dụng trên toàn Business hoặc theo permission cụ thể.", "Must", "Draft")
add_requirement("FR-PERM-003", "Store assignment", "User ngoài Owner chỉ được truy cập Store được gán hoặc được cấp quyền.", "Must", "Draft")
add_requirement("FR-PERM-004", "Action visibility", "Frontend chỉ hiển thị action mà user có quyền; khi quyền thay đổi, hệ thống phải refresh permission an toàn.", "Must", "Draft")
add_requirement("FR-PERM-005", "Billing privacy", "Store staff mặc định không được xem hóa đơn, payment method hoặc thông tin billing nhạy cảm.", "Must", "Confirmed")
add_requirement("FR-PERM-006", "No cross-store data", "User không được truy cập dữ liệu của Store không thuộc phạm vi quyền, kể cả bằng URL hoặc API trực tiếp.", "Must", "Draft")

# -----------------------
# 4. System Context
# -----------------------
add_section_header("4. System Context and Information Architecture")

doc.add_heading("4.1 Hierarchy", level=2)
add_callout("Business → Store → Fanpage")
add_bullets([
    "Business là phạm vi quản trị cao nhất.",
    "Store là đơn vị vận hành chính.",
    "Fanpage là kênh bán hàng thuộc Store.",
    "Product, Policy, Promotion, Assistant config và operational data bắt buộc thuộc Store.",
    "Members & Permissions và Usage & Billing nằm ở Business level.",
])

doc.add_heading("4.2 Navigation Model", level=2)
add_requirement("FR-NAV-001", "Dual sidebar", "Hệ thống phải sử dụng global navigation và Store-context navigation. Global navigation luôn tồn tại; Store sidebar chỉ xuất hiện khi user đang ở một Store.", "Must", "Confirmed")
add_requirement("FR-NAV-002", "Business dashboard after login", "Nếu user đã có Store, sau login hệ thống phải mở Business Dashboard. Nếu chưa có Store, hệ thống phải mở onboarding tạo Store.", "Must", "Confirmed")
add_requirement("FR-NAV-003", "Store switcher location", "Store switcher phải nằm trên header của main content, không đặt trong Store sidebar.", "Must", "Confirmed")
add_requirement("FR-NAV-004", "Keep current module", "Khi đổi Store, hệ thống nên giữ module hiện tại nếu user có quyền truy cập module đó ở Store mới.", "Should", "Confirmed")
add_requirement("FR-NAV-005", "Permission fallback", "Nếu user không có quyền truy cập module tương ứng ở Store mới, hệ thống phải chuyển về Store Dashboard và hiển thị thông báo ngắn.", "Must", "Confirmed")
add_requirement("FR-NAV-006", "Route context", "Mọi màn Store-level phải có store_id trong route hoặc context được xác thực.", "Must", "Confirmed")
add_requirement("FR-NAV-007", "Fanpage filter", "Các màn hỗ trợ Fanpage phải cho phép page_id=all hoặc một page_id cụ thể; Fanpage không tạo sidebar level 3.", "Must", "Confirmed")
add_requirement("FR-NAV-008", "Date range", "Các màn số liệu phải hỗ trợ date range phù hợp với phạm vi dữ liệu.", "Must", "Confirmed")
add_requirement("FR-NAV-009", "Locale prefix", "Frontend phải thống nhất có hoặc không sử dụng locale prefix như /vi trước khi migration route.", "Must", "TBD")

doc.add_heading("4.3 Store Context Rules", level=2)
add_requirement("BR-CTX-001", "Mandatory Store context", "Products, Policies, Promotions, Pages, Assistant, AI Quality, Messages, Deals và Store Dashboard phải hoạt động trong đúng một Store context.", "Must", "Confirmed")
add_requirement("BR-CTX-002", "Reset stale data", "Khi user đổi Store, UI không được tiếp tục hiển thị dữ liệu của Store cũ trong khi chờ tải dữ liệu mới.", "Must", "Draft")
add_requirement("BR-CTX-003", "Page ownership", "Một page_id được dùng làm filter phải thuộc Store hiện tại và user phải có quyền truy cập Store đó.", "Must", "Draft")

# -----------------------
# 5. Functional Requirements
# -----------------------
add_section_header("5. Functional Requirements")

doc.add_heading("5.1 Authentication and Account", level=2)
add_requirement("FR-AUTH-001", "User registration", "Hệ thống phải cho phép người dùng đăng ký tài khoản bằng phương thức được hỗ trợ và xác thực dữ liệu đầu vào.", "Must", "Draft")
add_requirement("FR-AUTH-002", "Login", "Hệ thống phải xác thực user và thiết lập session/token an toàn.", "Must", "Draft")
add_requirement("FR-AUTH-003", "Password recovery", "Hệ thống phải hỗ trợ forgot password và reset password bằng token có thời hạn.", "Must", "Draft")
add_requirement("FR-AUTH-004", "Email verification", "Hệ thống nên hỗ trợ verify email trước khi cho phép các hành động nhạy cảm hoặc production use.", "Should", "Draft")
add_requirement("FR-AUTH-005", "Session expiry", "Khi session hết hạn, hệ thống phải yêu cầu đăng nhập lại và không để lộ dữ liệu cache nhạy cảm.", "Must", "Draft")

doc.add_heading("5.2 Onboarding", level=2)
add_requirement("FR-ONB-001", "First Store creation", "User chưa có Store phải được dẫn vào flow tạo Store đầu tiên.", "Must", "Confirmed")
add_requirement("FR-ONB-002", "Progressive onboarding", "Onboarding dự kiến gồm Create Store, Connect Pages, Import Products, Policies, Assistant, Test Bot và Finish.", "Should", "Draft")
add_requirement("FR-ONB-003", "Resume onboarding", "Hệ thống nên lưu trạng thái onboarding để user có thể tiếp tục sau khi gián đoạn.", "Should", "TBD")
add_requirement("FR-ONB-004", "Skip optional steps", "Các bước không bắt buộc phải cho phép skip nếu Product Owner xác nhận.", "Could", "TBD")

doc.add_heading("5.3 Business Dashboard", level=2)
add_requirement("FR-BIZD-001", "Business summary KPIs", "Business Dashboard phải hiển thị tổng Store, fanpage hoạt động, pending human và deal phát hiện trong khoảng thời gian mặc định.", "Must", "Confirmed")
add_requirement("FR-BIZD-002", "Store cards", "Mỗi Store card phải hiển thị tên, ngành hàng, số fanpage, tin nhắn, bot xử lý, pending human, deal và trạng thái vận hành.", "Must", "Confirmed")
add_requirement("FR-BIZD-003", "Open Store Dashboard", "User phải có thể mở Store Dashboard từ Store card.", "Must", "Confirmed")
add_requirement("FR-BIZD-004", "Recent activity", "Business Dashboard nên hiển thị hoạt động hoặc cảnh báo gần đây liên quan đến Store và Fanpage.", "Should", "Confirmed")
add_requirement("FR-BIZD-005", "Permission-aware metrics", "Mọi KPI tổng phải chỉ tính các Store và dữ liệu mà user được phép xem.", "Must", "Draft")

doc.add_heading("5.4 Store Dashboard", level=2)
add_requirement("FR-STD-001", "Four tabs", "Store Dashboard phải có bốn tab: Tổng quan, Hội thoại, Sản phẩm và Fanpage.", "Must", "Confirmed")
add_requirement("FR-STD-002", "Dashboard scope", "Dashboard phải hỗ trợ Store, Fanpage và date range theo context.", "Must", "Confirmed")
add_requirement("FR-STD-003", "Overview KPIs", "Tab Tổng quan phải ưu tiên Tin nhắn hôm nay, Bot xử lý, Pending human và Deal phát hiện.", "Must", "Confirmed")
add_requirement("FR-STD-004", "Conversation analytics", "Tab Hội thoại phải phân tích workload AI/human, follow-up, thời gian phản hồi và khung giờ cao điểm.", "Must", "Confirmed")
add_requirement("FR-STD-005", "Product analytics", "Tab Sản phẩm phải hiển thị nhu cầu, deal, conversion nếu có dữ liệu và dữ liệu sản phẩm còn thiếu.", "Must", "Confirmed")
add_requirement("FR-STD-006", "Fanpage comparison", "Tab Fanpage phải so sánh tin nhắn, bot xử lý, pending, deal, automation rate và trạng thái page.", "Must", "Confirmed")
add_requirement("FR-STD-007", "No duplicated modules", "Deal và AI Quality không được tạo thành tab riêng trong Dashboard; chúng là workspace riêng.", "Must", "Confirmed")
add_requirement("FR-STD-008", "KPI formula governance", "Mọi KPI phải có định nghĩa, công thức, nguồn dữ liệu và refresh rule trước khi release production.", "Must", "Draft")

doc.add_heading("5.5 Messages", level=2)
add_requirement("FR-MSG-001", "Store-scoped inbox", "Messages phải hoạt động trong một Store và có thể lọc theo tất cả Fanpage hoặc một Fanpage cụ thể.", "Must", "Confirmed")
add_requirement("FR-MSG-002", "Required filters", "Messages phải hỗ trợ filter theo Fanpage, trạng thái, thời gian và người phụ trách nếu có.", "Must", "Confirmed")
add_requirement("FR-MSG-003", "Conversation context", "Mỗi hội thoại phải hiển thị Fanpage, trạng thái AI, pending human, deal detected, người phụ trách và thời gian chờ phù hợp.", "Must", "Confirmed")
add_requirement("FR-MSG-004", "Message sending", "User có quyền phải có thể gửi tin nhắn và nhận trạng thái gửi thành công/thất bại.", "Must", "Draft")
add_requirement("FR-MSG-005", "Human takeover", "Hệ thống phải có cơ chế người thật nhận xử lý hoặc tiếp quản hội thoại từ AI.", "Must", "Draft — cần state machine")
add_requirement("FR-MSG-006", "AI on/off", "User có quyền phải có thể bật hoặc tắt AI cho phạm vi được hỗ trợ.", "Must", "Draft — cần xác nhận scope conversation/page")
add_requirement("FR-MSG-007", "Assignee", "Hệ thống nên cho phép gán một người phụ trách chính cho hội thoại.", "Should", "Draft")
add_requirement("FR-MSG-008", "Realtime update", "Inbox và conversation thread nên cập nhật gần thời gian thực khi có tin mới hoặc thay đổi trạng thái.", "Should", "Draft")
add_requirement("FR-MSG-009", "No cross-store conversation", "Khi đổi Store, conversation list và thread phải reset và chỉ tải dữ liệu Store mới.", "Must", "Draft")

doc.add_heading("5.6 Deals", level=2)
add_requirement("FR-DEAL-001", "Deal workspace", "Hệ thống phải có màn Deal cần xử lý tách khỏi Messages để xử lý cơ hội bán hàng.", "Must", "Confirmed")
add_requirement("FR-DEAL-002", "Deal filters", "Deals phải hỗ trợ filter theo Fanpage, trạng thái, thời gian, sản phẩm và người phụ trách.", "Must", "Confirmed")
add_requirement("FR-DEAL-003", "Deal statuses", "Hệ thống dự kiến hỗ trợ Cần xác nhận, Đang follow-up, Chờ kho xác nhận, Đã hẹn gọi lại, Đã chốt và Đã bỏ qua.", "Must", "Draft — cần chốt workflow")
add_requirement("FR-DEAL-004", "Deal actions", "User có quyền phải có thể mở hội thoại, nhận/assign, thay đổi trạng thái, xác nhận, follow-up hoặc bỏ qua deal.", "Must", "Confirmed")
add_requirement("FR-DEAL-005", "Detection criteria", "Tiêu chí tạo Deal detected phải được định nghĩa và version hóa trước production.", "Must", "TBD")
add_requirement("FR-DEAL-006", "Estimated value", "Giá trị ước tính không phải KPI hoặc cột bắt buộc trong Phase 2 baseline.", "Must", "Confirmed")

doc.add_heading("5.7 Products", level=2)
add_requirement("FR-PROD-001", "Store-scoped products", "Mọi Product phải thuộc một Store.", "Must", "Confirmed")
add_requirement("FR-PROD-002", "Product CRUD", "User có quyền phải có thể tạo, xem, sửa và xóa hoặc archive Product.", "Must", "Draft")
add_requirement("FR-PROD-003", "Product media", "Product có thể chứa media theo phạm vi dữ liệu được hỗ trợ.", "Should", "Draft")
add_requirement("FR-PROD-004", "Product completeness", "Hệ thống nên phát hiện Product thiếu mô tả, ảnh, giá hoặc tồn kho nếu các field này thuộc model.", "Should", "Draft")
add_requirement("FR-PROD-005", "AI usage", "AI chỉ được sử dụng Product thuộc Store hiện tại khi trả lời hội thoại của Store đó.", "Must", "Draft")

doc.add_heading("5.8 Policies", level=2)
add_requirement("FR-POL-001", "Store-scoped policies", "Mọi Policy phải thuộc một Store.", "Must", "Confirmed")
add_requirement("FR-POL-002", "Policy CRUD", "User có quyền phải có thể tạo, xem, sửa và xóa hoặc archive Policy.", "Must", "Draft")
add_requirement("FR-POL-003", "Policy input", "Hệ thống có thể hỗ trợ nhập nội dung và/hoặc upload tài liệu theo định dạng được cho phép.", "Must", "Draft")
add_requirement("FR-POL-004", "Processing states", "Policy upload phải có trạng thái Uploading, Processing/Parsing, Ready và Failed hoặc tương đương.", "Must", "Draft")
add_requirement("FR-POL-005", "Retry", "User có quyền nên có thể retry tài liệu xử lý thất bại.", "Should", "Draft")
add_requirement("FR-POL-006", "Knowledge isolation", "AI chỉ được truy xuất Policy thuộc Store hiện tại.", "Must", "Draft")

doc.add_heading("5.9 Promotions", level=2)
add_requirement("FR-PROMO-001", "Store-scoped promotions", "Mọi Promotion phải thuộc một Store.", "Must", "Confirmed")
add_requirement("FR-PROMO-002", "Promotion lifecycle", "Promotion nên hỗ trợ draft, scheduled, active, expired và disabled hoặc bộ trạng thái tương đương.", "Should", "TBD")
add_requirement("FR-PROMO-003", "AI awareness", "AI chỉ được sử dụng promotion active và hợp lệ theo thời gian của Store hiện tại.", "Must", "TBD")

doc.add_heading("5.10 Fanpage Connections", level=2)
add_requirement("FR-PAGE-001", "Multiple pages per Store", "Một Store phải có thể liên kết nhiều Fanpage.", "Must", "Confirmed")
add_requirement("FR-PAGE-002", "Page list data", "Mỗi Fanpage phải hiển thị tên, avatar, Page ID, token status, AI status, message/pending/deal summary và lần đồng bộ gần nhất khi dữ liệu có sẵn.", "Must", "Confirmed")
add_requirement("FR-PAGE-003", "Page actions", "User có quyền phải có thể liên kết, ngắt liên kết, đồng bộ lại, bật/tắt AI và mở Dashboard đã filter theo Page.", "Must", "Confirmed")
add_requirement("FR-PAGE-004", "Token health", "Hệ thống phải phát hiện và hiển thị token hết hạn, thiếu quyền hoặc lỗi kết nối.", "Must", "Draft")
add_requirement("FR-PAGE-005", "Unique connection", "Một Fanpage không được liên kết đồng thời vào nhiều Store nếu chưa có business rule khác được chốt.", "Must", "TBD")
add_requirement("FR-PAGE-006", "Permission compliance", "Luồng kết nối phải tuân thủ quyền Meta được phê duyệt và trạng thái App Review.", "Must", "Draft")

doc.add_heading("5.11 AI Assistant", level=2)
add_requirement("FR-AIA-001", "Store AI configuration", "Mỗi Store phải có AI Assistant configuration riêng.", "Must", "Confirmed")
add_requirement("FR-AIA-002", "Configurable fields", "Hệ thống dự kiến hỗ trợ tên Agent, tone/phong cách, fallback message và các cấu hình hành vi được chốt.", "Must", "Draft")
add_requirement("FR-AIA-003", "Playground", "User có quyền nên có thể test AI trong Playground trước khi áp dụng production.", "Should", "Confirmed")
add_requirement("FR-AIA-004", "Save and version", "Thay đổi cấu hình AI phải có trạng thái lưu rõ; versioning hoặc audit history cần được xác nhận.", "Should", "TBD")
add_requirement("FR-AIA-005", "Page override", "Cần xác nhận AI config có cho phép override theo Fanpage hay chỉ dùng mặc định theo Store.", "Must", "TBD")

doc.add_heading("5.12 AI Quality", level=2)
add_requirement("FR-AIQ-001", "Action workspace", "AI Quality phải là workspace xử lý lỗi, không chỉ là dashboard analytics.", "Must", "Confirmed")
add_requirement("FR-AIQ-002", "Fallback conversations", "Hệ thống phải liệt kê conversation fallback và hỗ trợ hành động bổ sung dữ liệu hoặc test lại.", "Must", "Confirmed")
add_requirement("FR-AIQ-003", "Bad responses", "Hệ thống nên cho phép đánh dấu hoặc quản lý các câu trả lời AI chưa tốt.", "Should", "Confirmed")
add_requirement("FR-AIQ-004", "Missing knowledge", "Hệ thống nên gợi ý Product, Policy, Promotion hoặc FAQ còn thiếu.", "Should", "Confirmed")
add_requirement("FR-AIQ-005", "Suggested training", "Hệ thống có thể nhóm câu hỏi tương tự và đề xuất training/skill.", "Could", "Confirmed")
add_requirement("FR-AIQ-006", "Feedback ownership", "Cần xác định role nào được đánh dấu bad response và xác nhận bản sửa.", "Must", "TBD")

doc.add_heading("5.13 Members and Permissions", level=2)
add_requirement("FR-MEM-001", "Business-level management", "Thành viên và phân quyền phải được quản lý tập trung ở Business level.", "Must", "Confirmed")
add_requirement("FR-MEM-002", "Store assignments", "Hệ thống phải cho phép gán thành viên vào một hoặc nhiều Store theo role/permission được chốt.", "Must", "Draft")
add_requirement("FR-MEM-003", "Invitation lifecycle", "Hệ thống nên hỗ trợ lời mời Pending, Accepted, Expired và Revoked hoặc tương đương.", "Should", "TBD")
add_requirement("FR-MEM-004", "Owner protection", "Không được xóa hoặc hạ quyền Owner cuối cùng nếu chưa có cơ chế chuyển ownership.", "Must", "Draft")

doc.add_heading("5.14 Usage and Billing", level=2)
add_requirement("FR-BILL-001", "Business-level billing", "Usage & Billing phải nằm ở Business level.", "Must", "Confirmed")
add_requirement("FR-BILL-002", "Plan and quota", "Hệ thống phải hiển thị current plan, quota, current cycle và usage theo loại tài nguyên được chốt.", "Must", "Confirmed")
add_requirement("FR-BILL-003", "Store breakdown", "Usage phải có breakdown theo Store và có thể drill-down Fanpage khi cần.", "Must", "Confirmed")
add_requirement("FR-BILL-004", "Invoices and payment", "Owner hoặc role có quyền phải có thể xem invoice history và payment status.", "Must", "Confirmed")
add_requirement("FR-BILL-005", "Upgrade/downgrade", "Cơ chế upgrade/downgrade phải được xác định theo billing provider và policy thương mại.", "Must", "TBD")
add_requirement("FR-BILL-006", "Usage warning", "Hệ thống nên cảnh báo khi usage gần chạm quota hoặc có mức tiêu thụ bất thường.", "Should", "Confirmed")

doc.add_heading("5.15 Settings and Help", level=2)
add_requirement("FR-SET-001", "Business Settings", "Business Settings phải quản lý thông tin và cấu hình thuộc Business.", "Must", "Draft")
add_requirement("FR-SET-002", "Store Settings", "Store Settings phải quản lý thông tin và cấu hình thuộc Store.", "Must", "Confirmed")
add_requirement("FR-SET-003", "Personal Settings", "Cài đặt cá nhân phải truy cập từ avatar/header thay vì menu chính.", "Must", "Confirmed")
add_requirement("FR-HELP-001", "Help center", "Global navigation có thể cung cấp Trung tâm trợ giúp hoặc đường dẫn hỗ trợ.", "Should", "Confirmed")

# -----------------------
# 6. Data and Business Rules
# -----------------------
add_section_header("6. Data and Business Rules")

doc.add_heading("6.1 Core Entity Relationships", level=2)
t = doc.add_table(rows=1, cols=4)
for i, h in enumerate(["Entity", "Parent Scope", "Key Relationship", "Status"]):
    t.rows[0].cells[i].text = h
data_rows = [
    ("Business", "Account/Owner", "1 Business contains many Stores", "Draft — multi-Business TBD"),
    ("Store", "Business", "1 Store contains Products, Policies, Pages, Assistant config, Messages and Deals", "Confirmed"),
    ("Fanpage", "Store", "Many Fanpages may belong to one Store", "Confirmed"),
    ("Product", "Store", "Product belongs to exactly one Store", "Confirmed"),
    ("Policy", "Store", "Policy belongs to exactly one Store", "Confirmed"),
    ("Promotion", "Store", "Promotion belongs to exactly one Store", "Confirmed"),
    ("Conversation", "Store + Fanpage", "Conversation belongs to one Store and one Fanpage", "Draft"),
    ("Deal", "Store + Fanpage + Conversation", "Deal originates from conversation and may reference products", "Draft"),
    ("Member Assignment", "Business + Store", "User may have role/permission across assigned Stores", "Draft"),
]
for rd in data_rows:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.0)

doc.add_heading("6.2 System Business Rules", level=2)
rules = [
    ("BR-SYS-001", "Business is the highest management scope."),
    ("BR-SYS-002", "Store is the primary operational scope."),
    ("BR-SYS-003", "Fanpage is a Store-owned channel and filter, not a navigation level."),
    ("BR-SYS-004", "Billing and centralized member management belong to Business level."),
    ("BR-SYS-005", "Products, Policies, Promotions and AI configuration belong to Store."),
    ("BR-SYS-006", "Dashboard data must always identify its Business/Store/Fanpage/date scope."),
    ("BR-SYS-007", "Dashboard, Deals and AI Quality must not duplicate responsibilities."),
    ("BR-SYS-008", "No UI or API may expose cross-Store data without explicit permission."),
    ("BR-SYS-009", "Semantic theme tokens must be used instead of hard-coded brand color."),
]
for rid, text in rules:
    add_requirement(rid, "Business rule", text, "Must", "Confirmed" if rid not in ["BR-SYS-008"] else "Draft")

doc.add_heading("6.3 KPI Governance", level=2)
add_bullets([
    "Mỗi KPI phải có ID, tên, business meaning, công thức, data source, scope, refresh frequency và owner.",
    "Không release KPI mock như dữ liệu thật.",
    "Phải quy định timezone cho việc tính ngày/hôm nay.",
    "KPI tổng phải tôn trọng permission và Store scope.",
    "Nếu dữ liệu chưa đủ đáng tin cậy, UI phải hiển thị trạng thái unavailable/insufficient data thay vì số giả.",
])

# -----------------------
# 7. Interfaces
# -----------------------
add_section_header("7. External Interface and Integration Requirements")

doc.add_heading("7.1 Meta / Facebook", level=2)
add_requirement("INT-META-001", "Facebook Login", "Hệ thống sử dụng Facebook Login hoặc flow tương đương cho các tác vụ được chốt.", "Must", "Draft")
add_requirement("INT-META-002", "Page permissions", "Các quyền Meta phải phù hợp với phạm vi App Review được duyệt.", "Must", "Draft")
add_requirement("INT-META-003", "Messenger webhook", "Hệ thống phải nhận, deduplicate và xử lý webhook message theo cơ chế chịu retry.", "Must", "Draft")
add_requirement("INT-META-004", "Token lifecycle", "Hệ thống phải lưu, kiểm tra và xử lý token hết hạn/thiếu quyền một cách an toàn.", "Must", "Draft")
add_requirement("INT-META-005", "Reconnect", "UI phải cung cấp flow reconnect khi token hoặc permission không còn hợp lệ.", "Must", "Draft")

doc.add_heading("7.2 AI/LLM and Knowledge", level=2)
add_requirement("INT-AI-001", "LLM provider", "Hệ thống phải tích hợp LLM/AI provider qua lớp abstraction phù hợp với kiến trúc được chốt.", "Must", "Draft")
add_requirement("INT-AI-002", "Knowledge retrieval", "AI retrieval phải giới hạn theo Store và chỉ dùng nguồn dữ liệu hợp lệ.", "Must", "Draft")
add_requirement("INT-AI-003", "Fallback and safety", "Hệ thống phải có fallback behavior khi AI không đủ thông tin, timeout hoặc lỗi provider.", "Must", "Draft")
add_requirement("INT-AI-004", "Traceability", "AI response nên lưu thông tin cần thiết để truy vết nguồn dữ liệu, latency và lỗi theo chính sách privacy.", "Should", "TBD")
add_requirement("INT-AI-005", "Cost tracking", "Usage AI cần có khả năng được đo lường để phục vụ quota/billing nếu pricing model yêu cầu.", "Must", "Draft")

doc.add_heading("7.3 Storage and Document Processing", level=2)
add_requirement("INT-STO-001", "Media storage", "Media và tài liệu phải được lưu trên storage được cấu hình và truy cập qua quyền phù hợp.", "Must", "Draft")
add_requirement("INT-STO-002", "Document parser", "Policy upload phải được parse/index qua pipeline có trạng thái và lỗi truy vết được.", "Must", "Draft")
add_requirement("INT-STO-003", "File constraints", "Định dạng, dung lượng và số lượng file tối đa phải được xác nhận và hiển thị cho user.", "Must", "TBD")

doc.add_heading("7.4 Billing Provider", level=2)
add_requirement("INT-BILL-001", "Billing integration", "Nếu sử dụng payment/billing provider, hệ thống phải xử lý subscription, invoice và webhook một cách idempotent.", "Must", "TBD")
add_requirement("INT-BILL-002", "Payment status", "Payment status hiển thị phải đồng bộ với nguồn billing chính.", "Must", "TBD")

doc.add_heading("7.5 Future Channels", level=2)
add_callout(
    "NovaBot định hướng đa kênh, nhưng cần xác nhận Phase 2 production có chỉ triển khai Facebook/Messenger hay bao gồm Zalo, Instagram và TikTok. "
    "Các kênh chưa xác nhận không được xem là committed scope."
, fill="FEF3C7", border="D97706")

# -----------------------
# 8. NFR
# -----------------------
add_section_header("8. Non-functional Requirements")

doc.add_heading("8.1 Performance", level=2)
add_requirement("NFR-PERF-001", "Interactive response", "Các thao tác UI phổ biến nên phản hồi trạng thái loading trong vòng 300 ms và hoàn thành theo SLA được chốt.", "Should", "TBD")
add_requirement("NFR-PERF-002", "List pagination", "Các danh sách lớn phải hỗ trợ pagination, cursor hoặc virtualization phù hợp.", "Must", "Draft")
add_requirement("NFR-PERF-003", "Dashboard query", "Dashboard query phải có giới hạn date range và được tối ưu để không ảnh hưởng luồng message realtime.", "Must", "Draft")
add_requirement("NFR-PERF-004", "Realtime latency", "Mục tiêu latency cho message và status update cần được chốt.", "Must", "TBD")

doc.add_heading("8.2 Availability and Reliability", level=2)
add_requirement("NFR-REL-001", "Graceful degradation", "Khi một dịch vụ phụ trợ lỗi, hệ thống phải hiển thị lỗi có thể hiểu và tránh làm hỏng toàn bộ app shell.", "Must", "Draft")
add_requirement("NFR-REL-002", "Idempotency", "Webhook, payment callback và các action dễ retry phải được xử lý idempotent.", "Must", "Draft")
add_requirement("NFR-REL-003", "Data consistency", "Trạng thái hội thoại, deal, assignment và usage phải có source of truth rõ.", "Must", "Draft")
add_requirement("NFR-REL-004", "Backup and recovery", "RPO, RTO và chính sách backup phải được DevOps/Backend xác nhận trước production.", "Must", "TBD")

doc.add_heading("8.3 Scalability", level=2)
add_requirement("NFR-SCALE-001", "Multi-tenant isolation", "Kiến trúc phải hỗ trợ nhiều Business/Store mà không trộn dữ liệu.", "Must", "Draft")
add_requirement("NFR-SCALE-002", "Message burst", "Hệ thống phải xử lý burst webhook và tránh mất/nhân đôi message ngoài ngưỡng chấp nhận.", "Must", "Draft")
add_requirement("NFR-SCALE-003", "Background jobs", "Parsing, indexing, analytics và retry nên được xử lý bất đồng bộ khi phù hợp.", "Should", "Draft")

doc.add_heading("8.4 Usability and Accessibility", level=2)
add_requirement("NFR-UX-001", "Context clarity", "User phải luôn nhận biết đang ở Business, Store nào và Fanpage filter nào.", "Must", "Confirmed")
add_requirement("NFR-UX-002", "Responsive", "Ứng dụng phải hỗ trợ desktop và các viewport được chốt; behavior tablet/mobile cần xác nhận.", "Must", "TBD")
add_requirement("NFR-UX-003", "Keyboard and focus", "Các control chính phải có focus state và tương tác bàn phím phù hợp.", "Should", "Draft")
add_requirement("NFR-UX-004", "Color semantics", "Không được dùng màu duy nhất để truyền đạt trạng thái quan trọng.", "Should", "Draft")
add_requirement("NFR-UX-005", "Theme tokens", "Frontend phải dùng semantic theme token; không hard-code một accent cụ thể.", "Must", "Confirmed")

doc.add_heading("8.5 Compatibility", level=2)
add_requirement("NFR-COMP-001", "Browser support", "Danh sách browser/version hỗ trợ phải được chốt trước production.", "Must", "TBD")
add_requirement("NFR-COMP-002", "Locale and timezone", "Ngôn ngữ mặc định, locale route và timezone tính dữ liệu phải được xác nhận.", "Must", "TBD")

# -----------------------
# 9. States
# -----------------------
add_section_header("9. Error Handling and Common UI States")

doc.add_heading("9.1 Required Common States", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["State", "Requirement", "Notes"]):
    t.rows[0].cells[i].text = h
state_rows = [
    ("Initial loading", "Không hiển thị dữ liệu cũ hoặc dữ liệu sai scope.", "Skeleton/spinner theo Design System"),
    ("Refresh loading", "Giữ layout ổn định và báo đang cập nhật.", "Không khóa toàn màn nếu không cần"),
    ("Empty - no data", "Giải thích chưa có dữ liệu và CTA phù hợp.", "Khác search/filter empty"),
    ("Empty - search/filter", "Nêu không có kết quả và cho phép reset filter.", "Không dùng CTA tạo mới sai ngữ cảnh"),
    ("Error", "Thông báo dễ hiểu, có retry khi phù hợp.", "Không lộ stack trace"),
    ("401", "Yêu cầu đăng nhập lại.", "Xử lý session an toàn"),
    ("403", "Không có quyền; không hiển thị dữ liệu nhạy cảm.", "Có thể redirect hoặc no-access state"),
    ("404", "Resource/Store không tồn tại hoặc không còn truy cập.", "Điều hướng an toàn"),
    ("409", "Hiển thị conflict có thể hành động.", "Ví dụ duplicate page"),
    ("Validation", "Hiển thị lỗi tại field và summary khi cần.", "Giữ dữ liệu user nhập"),
    ("Processing", "Hiển thị trạng thái async và không cho action trùng.", "Upload/parse/index"),
]
for rd in state_rows:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.2)

doc.add_heading("9.2 Error Contract", level=2)
add_requirement("NFR-ERR-001", "Consistent error format", "API nên trả error code, user-safe message, trace/request ID và field errors khi phù hợp.", "Must", "Draft")
add_requirement("NFR-ERR-002", "Retry safety", "Retry không được tạo duplicate đối với các operation cần idempotency.", "Must", "Draft")
add_requirement("NFR-ERR-003", "User-safe messaging", "UI không hiển thị lỗi kỹ thuật hoặc thông tin nhạy cảm trực tiếp cho user.", "Must", "Draft")

# -----------------------
# 10. Security/privacy
# -----------------------
add_section_header("10. Security, Privacy and Authorization")

add_requirement("NFR-SEC-001", "Transport security", "Mọi traffic production phải sử dụng HTTPS/TLS.", "Must", "Draft")
add_requirement("NFR-SEC-002", "Secrets management", "Token, API key và credentials không được hard-code hoặc lưu ở client.", "Must", "Draft")
add_requirement("NFR-SEC-003", "Tenant isolation", "Mọi query nhạy cảm phải ràng buộc Business/Store ownership và permission.", "Must", "Draft")
add_requirement("NFR-SEC-004", "Input validation", "Backend phải validate input và file upload; không tin dữ liệu từ client.", "Must", "Draft")
add_requirement("NFR-SEC-005", "Sensitive data", "PII, message content và token phải được lưu/hiển thị theo chính sách privacy và least privilege.", "Must", "TBD")
add_requirement("NFR-SEC-006", "Audit security actions", "Các action như đổi quyền, kết nối page, bật/tắt AI, xóa dữ liệu và billing nên có audit log.", "Should", "Draft")
add_requirement("NFR-SEC-007", "Data retention", "Thời gian lưu message, conversation, audit log, file và AI trace phải được xác nhận.", "Must", "TBD")
add_requirement("NFR-SEC-008", "Deletion", "Cần xác định quy trình xóa Business, Store, Fanpage connection và dữ liệu liên quan.", "Must", "TBD")

# -----------------------
# 11. Logging
# -----------------------
add_section_header("11. Logging, Monitoring and Auditability")

add_requirement("NFR-OBS-001", "Request tracing", "API và background job nên có request/correlation ID để truy vết lỗi.", "Should", "Draft")
add_requirement("NFR-OBS-002", "Structured logs", "Log production nên có cấu trúc và không chứa secret hoặc PII không cần thiết.", "Must", "Draft")
add_requirement("NFR-OBS-003", "Webhook monitoring", "Phải theo dõi webhook failure, duplicate, backlog và retry.", "Must", "Draft")
add_requirement("NFR-OBS-004", "AI monitoring", "Nên theo dõi latency, error, fallback, provider usage và cost.", "Should", "Draft")
add_requirement("NFR-OBS-005", "Alerting", "Ngưỡng cảnh báo cho lỗi nghiêm trọng, token hết hạn, queue backlog và quota phải được xác nhận.", "Must", "TBD")
add_requirement("NFR-OBS-006", "Audit log", "Audit log phải ghi actor, action, target, timestamp, result và context phù hợp.", "Should", "Draft")

# -----------------------
# 12. Deployment migration
# -----------------------
add_section_header("12. Deployment, Migration and Backward Compatibility")

doc.add_heading("12.1 Phase 1 → Phase 2 Migration", level=2)
add_bullets([
    "Audit route, component, API và data model trước khi rebuild.",
    "Xác định mỗi màn là Keep, Restyle, Refactor, Move, Replace, Remove/Redirect hoặc New Build.",
    "App Shell và Store context phải được triển khai trước khi migrate module.",
    "Route cũ phải redirect hoặc bị deprecate có kiểm soát.",
    "Không xóa dữ liệu Phase 1 nếu chưa có migration plan và backup.",
])

add_requirement("MIG-SYS-001", "Backward-compatible APIs", "API dùng chung nên được version hoặc duy trì tương thích trong thời gian migration nếu FE cũ và mới chạy song song.", "Should", "TBD")
add_requirement("MIG-SYS-002", "Route redirects", "Route analytics, billing và team cũ phải redirect đến route mới theo mapping đã chốt.", "Must", "Confirmed")
add_requirement("MIG-SYS-003", "Feature flags", "Có thể sử dụng feature flag để rollout app shell/module mới theo môi trường hoặc nhóm user.", "Should", "TBD")
add_requirement("MIG-SYS-004", "Rollback", "Mỗi release migration cần có rollback plan phù hợp với thay đổi code và schema.", "Must", "Draft")
add_requirement("MIG-SYS-005", "Mock data removal", "Không được release production với KPI mock được trình bày như dữ liệu thật.", "Must", "Draft")

doc.add_heading("12.2 Environment Requirements", level=2)
add_requirement("ENV-001", "Environment separation", "Development, staging và production phải được tách cấu hình và credentials.", "Must", "Draft")
add_requirement("ENV-002", "Staging UAT", "Các flow quan trọng phải được UAT trên staging trước production.", "Must", "Draft")
add_requirement("ENV-003", "Database migration", "Schema migration phải có version, review và rollback/forward strategy.", "Must", "Draft")

# -----------------------
# 13. Traceability
# -----------------------
add_section_header("13. Requirement Traceability and Change Control")

doc.add_heading("13.1 Traceability Model", level=2)
add_callout("SRS Requirement → Screen Specification → Feature / Acceptance Criteria → Jira Story → Test Case → Release")
add_bullets([
    "Mỗi requirement có ID duy nhất.",
    "Screen Specification phải tham chiếu requirement cấp hệ thống liên quan.",
    "Jira Story phải link về Screen Specification và copy AC cần thiết.",
    "QA test case phải tham chiếu Story/AC.",
    "Thay đổi requirement đã Approved phải ghi Decision Log và impact.",
])

doc.add_heading("13.2 Priority Convention", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Priority", "Meaning", "Release rule"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("Must", "Bắt buộc để flow/safety/scope hoạt động đúng", "Không release nếu chưa đạt"),
    ("Should", "Giá trị cao nhưng có thể defer có chấp thuận", "Cần quyết định rõ"),
    ("Could", "Nice-to-have hoặc optimization", "Không block release"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

doc.add_heading("13.3 Change Control", level=2)
add_bullets([
    "Product Owner là người chốt scope và business rule.",
    "Designer, FE, BE/AI và QA có quyền raise risk hoặc đề xuất thay đổi.",
    "Mọi thay đổi ảnh hưởng route, permission, data model, KPI hoặc release scope phải ghi Decision Log.",
    "Story đang In Progress chỉ thay đổi scope sau khi đánh giá impact và re-estimate.",
])

# -----------------------
# 14. Assumptions
# -----------------------
add_section_header("14. Assumptions and Constraints")

doc.add_heading("14.1 Current Assumptions", level=2)
add_bullets([
    "NovaBot Phase 2 tiếp tục sử dụng mô hình một Business có nhiều Store.",
    "Một Store có thể có nhiều Fanpage.",
    "Phase 1 code/API sẽ được tái sử dụng khi phù hợp thay vì rebuild toàn bộ.",
    "Business Dashboard là màn mặc định sau login khi đã có Store.",
    "Light theme và semantic token là baseline.",
    "Facebook/Messenger là kênh đã có nền tảng rõ nhất ở thời điểm lập SRS.",
])

doc.add_heading("14.2 Known Constraints", level=2)
add_bullets([
    "Role model và permission chi tiết chưa được chốt cuối.",
    "KPI formula và dữ liệu thật cho Dashboard chưa hoàn chỉnh.",
    "Billing provider và pricing model chưa được xác nhận.",
    "Meta App Review và quyền thực tế có thể giới hạn một số flow.",
    "Một số route Phase 1 và API status cần FE/BE audit.",
    "SLA hiệu năng, availability, backup và retention chưa có con số chính thức.",
])

# -----------------------
# 15. Open Questions
# -----------------------
add_section_header("15. Open Questions / Confirmations Required")
add_callout(
    "Các câu hỏi dưới đây không chặn việc bắt đầu Screen Specification cho App Shell, nhưng cần được chốt dần để SRS chuyển từ v0.9 Draft sang v1.0 Approved.",
    fill="FEF3C7", border="D97706"
)

questions = [
    ("Q-001", "Role model chính thức là 3 role Owner/Admin/Staff hay mô hình chi tiết Business Owner/Business Admin/Store Manager/Sales/Support?"),
    ("Q-002", "Một user/account có thể sở hữu hoặc tham gia nhiều Business hay chỉ một Business?"),
    ("Q-003", "Business entity đã tồn tại trong database/backend hay cần tạo mới trong Phase 2?"),
    ("Q-004", "Permission dùng fixed role hay permission động theo từng feature/Store?"),
    ("Q-005", "Phạm vi kênh production Phase 2: chỉ Facebook/Messenger hay gồm Zalo, Instagram, TikTok?"),
    ("Q-006", "Human takeover state machine cụ thể: AI tự tắt khi nào, ai được bật lại, phạm vi conversation hay Page?"),
    ("Q-007", "Tiêu chí Pending human và SLA/thời gian chờ được tính như thế nào?"),
    ("Q-008", "Tiêu chí Deal detected, Deal confirmed và conversion là gì?"),
    ("Q-009", "AI Assistant có config mặc định theo Store và override theo Fanpage không?"),
    ("Q-010", "Product model production gồm field bắt buộc nào: SKU, price, stock, category, media, variants?"),
    ("Q-011", "Policy hỗ trợ định dạng file nào, dung lượng tối đa và lifecycle xóa/indexing ra sao?"),
    ("Q-012", "Promotion có nằm trong release đầu Phase 2 không? Các trạng thái và rule thời gian?"),
    ("Q-013", "Billing model tính theo messages, AI replies, conversations, page hay combination?"),
    ("Q-014", "Billing provider/payment gateway nào sẽ dùng?"),
    ("Q-015", "KPI Dictionary: công thức chính xác cho Bot xử lý, Pending human, Deal phát hiện và automation rate?"),
    ("Q-016", "Timezone mặc định của Business/Store và quy tắc tính 'Hôm nay'?"),
    ("Q-017", "Ngôn ngữ, locale prefix và khả năng multi-language của app?"),
    ("Q-018", "Browser/device support: desktop-only trong Phase 2 hay cần tablet/mobile responsive đầy đủ?"),
    ("Q-019", "Performance SLA: page load, API p95, message realtime latency và Dashboard query?"),
    ("Q-020", "Availability, backup RPO/RTO và disaster recovery target?"),
    ("Q-021", "Data retention và deletion policy cho messages, files, AI traces và audit logs?"),
    ("Q-022", "Notification scope: in-app, email, Slack/Zalo, browser push?"),
    ("Q-023", "Audit log có phải Must-have trong release đầu không? Ai được xem?"),
    ("Q-024", "Onboarding step nào bắt buộc, step nào được skip, có resume không?"),
    ("Q-025", "Route production có prefix /vi hoặc cơ chế i18n nào?"),
]
t = doc.add_table(rows=1, cols=4)
for i, h in enumerate(["ID", "Question / Decision", "Owner", "Status"]):
    t.rows[0].cells[i].text = h
for qid, question in questions:
    cells = t.add_row().cells
    cells[0].text = qid
    cells[1].text = question
    cells[2].text = "Product Owner / Team"
    cells[3].text = "Open"
style_table(t, first_col_bold=True, font_size=7.8)

# -----------------------
# Appendix A
# -----------------------
doc.add_page_break()
doc.add_heading("Appendix A. Route Map", level=1)
route_groups = [
    ("Auth", ["/login", "/register", "/forgot-password", "/reset-password", "/verify-email"]),
    ("Onboarding", ["/onboarding/create-store", "/onboarding/connect-pages", "/onboarding/import-products", "/onboarding/policies", "/onboarding/assistant", "/onboarding/test-bot", "/onboarding/finish"]),
    ("Business", ["/business/dashboard", "/business/members", "/business/usage-billing", "/business/settings", "/help"]),
    ("Store", [
        "/stores/{store_id}/dashboard",
        "/stores/{store_id}/dashboard/conversations",
        "/stores/{store_id}/dashboard/products",
        "/stores/{store_id}/dashboard/fanpages",
        "/stores/{store_id}/messages",
        "/stores/{store_id}/deals",
        "/stores/{store_id}/products",
        "/stores/{store_id}/products/{product_id}",
        "/stores/{store_id}/policies",
        "/stores/{store_id}/promotions",
        "/stores/{store_id}/pages",
        "/stores/{store_id}/assistant",
        "/stores/{store_id}/assistant/playground",
        "/stores/{store_id}/ai-quality",
        "/stores/{store_id}/training",
        "/stores/{store_id}/skills",
        "/stores/{store_id}/settings",
    ]),
]
for group, routes in route_groups:
    doc.add_heading(group, level=2)
    add_bullets(routes)

doc.add_heading("Redirects / Deprecated Routes", level=2)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Old route", "New route", "Reason"]):
    t.rows[0].cells[i].text = h
for rd in [
    ("/stores/{store_id}/analytics", "/stores/{store_id}/dashboard", "Analytics gộp vào Dashboard"),
    ("/stores/{store_id}/billing", "/business/usage-billing?store_id=...", "Billing chuyển lên Business"),
    ("/stores/{store_id}/team", "/business/members?store_id=...", "Members chuyển lên Business"),
]:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.3)

# Appendix B
doc.add_heading("Appendix B. Requirement ID Convention", level=1)
t = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["Prefix", "Meaning", "Example"]):
    t.rows[0].cells[i].text = h
id_rows = [
    ("FR", "Functional Requirement", "FR-MSG-001"),
    ("BR", "Business Rule", "BR-SYS-001"),
    ("NFR", "Non-functional Requirement", "NFR-SEC-001"),
    ("INT", "Integration Requirement", "INT-META-001"),
    ("MIG", "Migration Requirement", "MIG-SYS-001"),
    ("ENV", "Environment Requirement", "ENV-001"),
    ("Q", "Open Question", "Q-001"),
]
for rd in id_rows:
    cells = t.add_row().cells
    for i, val in enumerate(rd):
        cells[i].text = val
style_table(t, first_col_bold=True, font_size=8.5)

# Appendix C
doc.add_heading("Appendix C. System-level Release Acceptance", level=1)
add_bullets([
    "Requirement Must trong release scope có status Approved và có traceability tới Story/Test.",
    "Permission và tenant isolation được test ở Backend.",
    "Không có critical cross-Store data leakage.",
    "Không có KPI mock được hiển thị như dữ liệu thật.",
    "Route migration và redirect đã test.",
    "Loading, empty, error, 401/403/404 states của critical flow đã được kiểm tra.",
    "Meta token/webhook failure có monitoring và recovery phù hợp.",
    "Backup/migration/rollback plan được duyệt cho release có schema change.",
    "QA pass và Product Owner sign-off trên staging.",
])

add_callout(
    "Next document: Screen Specification cho App Shell & Navigation. Sau đó mới tạo Jira Epic/Story cho Sprint 1."
)

doc.save(output_path)
print(f"Created: {output_path}")
